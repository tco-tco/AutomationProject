
from belgenetUser import username, password
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.select import Select
from docx import Document

from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from CED_TCO import ExcelDosya
import konuCharHarf


import os
import shutil
import time


class Belgenet:
    def __init__(self, username, password):
        self.browser=webdriver.Chrome("C:/Users/tugbacanan.oguz/Desktop/python_temelleri_iş/PROJECT/chromedriver.exe")
        #self.browser=webdriver.Safari()
        self.username=username
        self.password=password
        
             
    
    def signIn(self):
        self.browser.get("https://e-belge.tarim.gov.tr/edys-web/sistemeGiris.xhtml")
        self.browser.implicitly_wait(60)

        self.browser.find_element(By.XPATH,"//*[@id='parolaSertifikaAccordion:uForm:txtUKullaniciAdi']").send_keys(self.username)
        self.browser.implicitly_wait(60)
        self.browser.find_element(By.XPATH,"//*[@id='loginUSifre']").send_keys(self.password)
        self.browser.implicitly_wait(60)

        self.browser.implicitly_wait(60)

        self.browser.find_element(By.XPATH,"/html/body/div[5]/div[2]/div/div[2]/div/div[2]/div[2]/div/form/button/span").click()

        self.browser.implicitly_wait(60)
#KİŞİSEL ARŞİV

        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//*[@id='leftMenuForm:kisiselArsivPanel:kisiselArsivKlasoruTab_header']"))).click() #kişisel arşiv
        
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"/html/body/div[9]/div[2]/form/div[2]/div[2]/div/ul/li/span/span[1]"))).click()  #altmenü kişisel arşiv

        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.CSS_SELECTOR,"#leftMenuForm\:kisiselArsivPanel\:kisiselArsivKlasoruTree\:0_1 > span:nth-child(1) > span:nth-child(2)"))).click()   # 2. sıradaki klasörü seçme
        
        self.browser.implicitly_wait(60)
        time.sleep(2)
        


    def letter(self):
        import time
        konuText=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[1]/td[2]/div[1]/h3").text   #yazının konusunu text halinde veriyor      
        self.browser.implicitly_wait(60)
        time.sleep(2)    
        print(konuText)
        konuListim=str(konuText).split()
        konuListim.remove("Konu:")
        for i in range(len(konuListim)):
            if konuListim[i]=="Iı.":
                konuListim[i]="II."
            elif konuListim[i]=="Iv.":
                konuListim[i]="IV."
            elif konuListim[i]=="Iı-":
                konuListim[i]="II-"
            elif konuListim[i]=="Iv-":
                konuListim[i]="IV-"
            elif konuListim[i]=="Iı":
                konuListim[i]="II"
            elif konuListim[i]=="Iv":
                konuListim[i]="IV"


        konu=" ".join(konuListim)
        if "Toplantı Tarihleri" in konuText:         
            joinKonuSon=str(konu.replace("Toplantı Tarihleri",""))
        elif "İDK Toplantısı" in konuText:         
            joinKonuSon=str(konu.replace("İDK Toplantısı",""))
        elif "Halkın Katılımı Toplantısı" in konuText:         
            joinKonuSon=str(konu.replace("Halkın Katılımı Toplantısı",""))
        else:
            joinKonuSon=str(konu)
        print(joinKonuSon)      #nihai konuyu verir       
        print("OKEYİM-konu")  

        sayi=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[3]/td[1]/div[1]").text   #yazının sayısını text halinde veriyor 
        self.browser.implicitly_wait(60)
        time.sleep(2)
        print(sayi)

        geldY=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[2]/td[1]/div[1]").text   #yazının geldiği yeri text halinde veriyor 
        self.browser.implicitly_wait(60)
        print(geldY)
        time.sleep(2)

        sayiYaziList=str(sayi).split("/")
        sayiYazi=(str(sayiYaziList[2]).split("-"))[-1].strip()
        print(sayiYazi)
        sayiYazim=str(sayiYaziList[2]).strip()
        print(sayiYazim)

        tarihYazi=str(sayiYaziList[1]).strip("Sayı: ").strip()
        print(tarihYazi)

# GELEN YAZI 

        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//*[@id='esm_246802192_emi_286053589']"))).click()      # gelen kutusunu tıklar
        self.browser.implicitly_wait(60)
        time.sleep(5)
        gelYaziNo=self.browser.find_element(By.CSS_SELECTOR,"a[id='esm_246802192_emi_286053589'] span[class='ui-menuitem-text']").text

        self.browser.implicitly_wait(60)
        time.sleep(2)        
        gelYaziNoList=str(gelYaziNo).lstrip("Gelen Evraklar ").split("/")
        gelYaziNom=int(gelYaziNoList[1].rstrip(")"))
        print(gelYaziNom)       
        

        sel100 = Select(self.browser.find_element(By.XPATH,"//div[@class='ui-paginator ui-paginator-top ui-widget-header']//select[@name='mainInboxForm:inboxDataTable_rppDD']"))   #sayfada görünen yazı sayısını 100'e çıkarır
        self.browser.implicitly_wait(60)
        time.sleep(3)

        sel100.select_by_visible_text("100")
        self.browser.implicitly_wait(60)
        time.sleep(4)
        print("OK-100")

        yazinom=""
        geldigiYer=""
        for i in range(0,gelYaziNom):

            gelenKonu=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[1]/td[2]/div[1]/h3").text    # tüm gelen yazıların konusunu alır
            self.browser.implicitly_wait(60)
            print(gelenKonu)

            geldigiYerr=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[2]/td/div").text        # tüm gelen yazıların geldiği yeri alır                                                        
            self.browser.implicitly_wait(60)
            geldigiYer=geldigiYerr
            print(geldigiYer)    

            gelenSayi=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[4]/td/div").text        # tüm gelen yazıların tarihini alır
            self.browser.implicitly_wait(60)
            print(gelenSayi)   
            gelenSayiList=str(gelenSayi).split("/")
            print(gelenSayiList)

            if "-" in str(gelenSayiList[1]):
                gelenSayison=(str(gelenSayiList[1]).split("-"))[-1]
                
            else:
                gelenSayison=str(gelenSayiList[1]).lstrip("No: ")
            
            print(gelenSayison)               

            gelenTarih=str(gelenSayiList[0]).strip("Evrak Tarihi: ").strip()
            print(gelenTarih)
            print(i)
              

          
            if (konuText == gelenKonu) and ((sayiYazi==gelenSayison) or (tarihYazi==gelenTarih)):
     
                self.browser.find_element(By.XPATH,f"//table[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']").click()  
                self.browser.implicitly_wait(60)
                time.sleep(1)                                  
                print("OK-gelen yazi bulundu")   
                print(i) 
                break  
                

    # FRAME YAZI ÖNİZLEME
        self.browser.switch_to.frame("ustYaziOnizlemeId1")  # en sağdaki frame'e iframe id'si ile gitme
        self.browser.implicitly_wait(60)
        time.sleep(2)
        print("OK2")
        textim=self.browser.find_elements(By.XPATH,"//*[@id='viewer']/div[1]/div[2]/div")
        self.browser.implicitly_wait(60)
        time.sleep(3)
        print("OK3")
        #print(textim[12].text)

        nihaiYaziIc=""
        cedMetinListem=[]
        for text in textim:
            cedMetinListem.append(text.text)
        print(cedMetinListem)
        geldYer1=""
        geldYer2=""
        pdfokunabilirmi=""
        if int(len(cedMetinListem))<4:
            okunmayanpdf="okunmayanpdf"
            pdfokunabilirmi=okunmayanpdf
            print("okunmayan pdf:(")

        elif int(len(cedMetinListem))>4:
            okunabilirpdf="okunabilirpdf"
            pdfokunabilirmi=okunabilirpdf

            yaziIcIndex=(cedMetinListem.index("DAĞITIM YERLERİNE"))+1
            print(yaziIcIndex)
            tcIndex=(cedMetinListem.index("T.C."))
            geldYer1=cedMetinListem[(int(tcIndex)+1)]
            geldYer2=cedMetinListem[(int(tcIndex)+2)]

            if cedMetinListem[yaziIcIndex] =="İlgi" or cedMetinListem[(int(yaziIcIndex)+1)]==":":
                
                print("yazının ilgisi var")
                yaziIcIndexim=(cedMetinListem.index("DAĞITIM YERLERİNE"))+4
                print(cedMetinListem[yaziIcIndex])
                print(cedMetinListem[(int(yaziIcIndex)+1)])
                yaziIcIndex=yaziIcIndexim
                print(yaziIcIndex)
                
            else:
                print("yazının ilgisi yok")
                print(cedMetinListem[yaziIcIndex])
            

                        
            yaziMetniListe=[]
            for xindex in range((yaziIcIndex+1),((len(cedMetinListem)-(yaziIcIndex+1)))):
                yaziMetni=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{xindex}]").text                         
                print(yaziMetni)            
                yaziMetniListe.append(yaziMetni)
            
            joinYaziMetni=" ".join(yaziMetniListe)
            #print(joinYaziMetni)
            splitYaziMetni=joinYaziMetni.split()

            for i in range(len(splitYaziMetni)):
                if splitYaziMetni[i]=="Iı.":
                    splitYaziMetni[i]="II."
                elif splitYaziMetni[i]=="Iv.":
                    splitYaziMetni[i]="IV."
                elif splitYaziMetni[i]=="Iı-":
                    splitYaziMetni[i]="II-"
                elif splitYaziMetni[i]=="Iv-":
                    splitYaziMetni[i]="IV-"
                elif splitYaziMetni[i]=="Iı":
                    splitYaziMetni[i]="II"
                elif splitYaziMetni[i]=="Iv":
                    splitYaziMetni[i]="IV"

            print(splitYaziMetni)

            indexProjesi=0
            if "projesi" in splitYaziMetni:            
                indexProje=int(splitYaziMetni.index("projesi"))+1
                
                indexProjesi=int(indexProje)
            elif "Projesi" in splitYaziMetni:            
                indexProje=int(splitYaziMetni.index("Projesi"))+1
                
                indexProjesi=int(indexProje)
            elif "projesiyle" in splitYaziMetni:            
                indexProje=int(splitYaziMetni.index("projesiyle"))+1
                
                indexProjesi=int(indexProje)
            else:
                indexProje=int(splitYaziMetni.index("ederim."))+1
                indexProjesi=int(indexProje)

            nihaiMetin=" ".join(splitYaziMetni[:(indexProjesi)])
            #print(nihaiMetin)        
            nihaiİl=nihaiMetin.replace("İli","ili")
            nihaiİlce=nihaiİl.replace("İlçesi","ilçesi")
            
            
            # nihai yazı metnini mevkiyi düzelterek verir
            while True:
                if "mevkiinde" in nihaiİlce:
                    nihaimevkii=nihaiİlce.replace("mevkiinde","mevkisinde")
                    nihaiYaziIci=nihaimevkii
                elif "Mevkiinde" in nihaiİlce:
                    nihaimevkii=nihaiİlce.replace("Mevkiinde","mevkisinde")
                    nihaiYaziIci=nihaimevkii
                elif "Mevkii'nde" in nihaiİlce:
                    nihaimevkii=nihaiİlce.replace("Mevkii'nde","mevkisinde")
                    nihaiYaziIci=nihaimevkii
                elif "Mevkii'" in nihaiİlce:
                    nihaimevkii=nihaiİlce.replace("Mevkii","mevkisi")
                    nihaiYaziIci=nihaimevkii
                elif "mevkii'" in nihaiİlce:
                    nihaimevkii=nihaiİlce.replace("mevkii","mevkisi")
                    nihaiYaziIci=nihaimevkii
                else:
                    nihaiYaziIci=nihaiİlce
                print(nihaiYaziIci)
                break
            nihaiYaziIc=nihaiYaziIci
        
        print(pdfokunabilirmi)
        
        import re
        listem=re.findall(r'\d+', konuText)    # yazının konusu içindeki sayıları extract ediyor
        print(listem)
        listemson=[]
        listem2=[]
        for i in listem:
            print(i)
            print(int(max(listem)))
            if int(max(listem))<2100:          
                listem.clear()
                listemson=listem
                #print(listemson)

            elif int(i)>2100:
                listem2.append(i)
                listemson=listem2
                #print(listemson)
        
        print(listemson)

        
        self.browser.execute_script("window.open()")
        self.browser.implicitly_wait(60)
        print(self.browser.window_handles)
        self.browser.switch_to.window(self.browser.window_handles[1])   # yeni pencere/sekme
        self.browser.implicitly_wait(60)
        self.browser.get("http://eced.csb.gov.tr/ced/jsp/portal/main2.htm") 
        self.browser.implicitly_wait(60)
        
        ale = self.browser.switch_to.alert  # ced sayfasında çıkan uyarıyı kaldırma
        ale.accept()
        # self.browser.find_element(By.XPATH,"//*[@id='details-button']").click()
        # self.browser.find_element(By.XPATH,"//*[@id='proceed-link']").click()

        self.browser.find_element(By.XPATH,"//*[@id='j_username']").send_keys("username")
        self.browser.implicitly_wait(60)
        self.browser.find_element(By.XPATH,"//*[@id='j_password']").send_keys("password")

        self.browser.implicitly_wait(60)

        self.browser.find_element(By.XPATH,"/html/body/header/nav/ul/div/form/div/span[5]/input").click()   # şifre giriş butonununa tıklama
        #self.browser.get("http://eced.csb.gov.tr/ced/jsp/disKurum/listProjeIdk.htm?ad=&ara=Ara")
        self.browser.implicitly_wait(60)

        
        if len(listemson)==0:           #yazının konusundaki numaraların atıldığı listenin boş olması halinde  
            print(nihaiYaziIc)
            print(joinKonuSon)
            while True:                     
                projeAdi=input("""
                

                ********************************
                Proje adını/bir kısmını giriniz: 
                ********************************
                
                
                
                """)
                self.browser.find_element(By.XPATH,"//*[@id='ad']").clear()
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='ad']").send_keys(projeAdi)  # kullanıcıdan alınan proje adını ced sisteminde arama
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='ara']").click()   
                self.browser.implicitly_wait(60)       

                kayit1=self.browser.find_element(By.XPATH,"//span[@class='pagebanner']").text #???????#2 kayıt eklemen lazım!!!!! 
                time.sleep(2)
                self.browser.implicitly_wait(60) 

                kayit2=self.browser.find_element(By.XPATH,"//div[@id='main']").text   #kayıt bulunamadı
                time.sleep(2)
                self.browser.implicitly_wait(60) 
                print(kayit1)
                print(kayit2)
                


                if kayit1=="Bir kayıt bulundu.":
                    self.browser.implicitly_wait(60)

                    self.browser.find_element(By.CSS_SELECTOR,".b > td:nth-child(12) > input:nth-child(1)").click()   #ced satırındaki "Ayrıntılar" butonu

                    self.browser.implicitly_wait(60) 
                    break                   
                elif kayit2=="Kayıt Bulunamadı":  
                                    
                    print(f"Kayıt yok- {projeAdi}") 

                elif "tane kayıt bulundu;" in kayit1:
                    cedSecimInput=input("""

                    
                    *****************************
                    Projeyi listeden sen seç(y/n): 
                    *****************************
                    
                    
                    
                    """)
                    print(cedSecimInput)
                    if cedSecimInput=="y":
                        print("OK- ÇED seçimi tamamlandı")
                        break
                    elif cedSecimInput=="n":
                        print("Seçim yapılmadı, proje adına geri dönülecek...")
                        continue
                elif "kayıt bulundu," in kayit1:
                    cedSecimInput=input("""
                    

                    *****************************
                    Projeyi listeden sen seç(y/n):
                    ***************************** 
                    
                    
                    
                    """)
                    print(cedSecimInput)
                    if cedSecimInput=="y":
                        print("OK- ÇED seçimi tamamlandı")
                        break
                    elif cedSecimInput=="n":
                        print("Seçim yapılmadı, proje adına geri dönülecek...")
                        continue
                    
                break       
            

                
                
                
        elif len(listemson)!=0:    #yazının konusundaki numaraların atıldığı listenin dolu olması halinde 
            
            for num in listemson:  
                print(num)

                print(joinKonuSon)
                print(listemson)
                print(i)



                self.browser.find_element(By.XPATH,"/html/body/div[1]/div[2]/div[1]/div/ul[1]/form/input[1]").clear()
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='ad']").send_keys(num)  # eğer yazının konusu içindeki numara uygunsa bu numarayı ced sisteminde aratır
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='ara']").click()
                self.browser.implicitly_wait(60)


                kayit2=self.browser.find_element(By.XPATH,"//div[@id='main']").text   #kayıt bulunamadı
                self.browser.implicitly_wait(60) 
                time.sleep(2)                
                #print(kayit1)
                print(f"kayit1 {kayit2}")
                print("x1")
                if "Bir kayıt bulundu" in kayit2:
                    print(f"kayit2 {kayit2}")
                    self.browser.find_element(By.CSS_SELECTOR,".b > td:nth-child(12) > input:nth-child(1)").click()   #ced satırındaki "Ayrıntılar" butonu
                    self.browser.implicitly_wait(60)

                    yazinom=num 
                    break 
                else:           
                    while True:
                        kayit2=self.browser.find_element(By.XPATH,"//div[@id='main']").text   #kayıt bulunamadı
                        self.browser.implicitly_wait(60) 
                        time.sleep(2)                
                        #print(kayit1)
                        print(f"kayit1 {kayit2}")
                        print("x1")
                        if "Bir kayıt bulundu" in kayit2:
                            print(f"kayit2 {kayit2}")

                            self.browser.implicitly_wait(60)
                            #self.browser.find_element(By.XPATH,"//*[@id='pr']/tbody/tr/td[12]/input").click()   #ced satırındaki "Ayrıntılar" butonu
                            self.browser.find_element(By.CSS_SELECTOR,".b > td:nth-child(12) > input:nth-child(1)").click()   #ced satırındaki "Ayrıntılar" butonu
                            #html body div.wrap div#content div#main table#pr.app tbody tr.b td input
                            self.browser.implicitly_wait(60)
                            yazinom=num 
                            break 
                        
                        elif "Kayıt Bulunamadı" in kayit2:                              
                            print(f"Kayıt yok") 
                            time.sleep(0.5)   
                            projeAdi=input("""
                            

                            ********************************
                            Proje adını/bir kısmını giriniz: 
                            ********************************
                            
                            
                            
                            """)
                            self.browser.find_element(By.XPATH,"//*[@id='ad']").clear()
                            self.browser.implicitly_wait(60)
                            self.browser.find_element(By.XPATH,"//*[@id='ad']").send_keys(projeAdi)  # eğer yazının konusu içindeki numara uygunsa bu numarayı ced sisteminde aratır
                            self.browser.implicitly_wait(60)
                            self.browser.find_element(By.XPATH,"//*[@id='ara']").click()
                            self.browser.implicitly_wait(60)    

                        
                        elif "tane kayıt bulundu;" in kayit2:
                            print( "Çok kayıt bulundu")
                            cedSecimInput=input("""
                            
                            
                            ******************************
                            Projeyi listeden sen seç(y/n): 
                            ******************************
                            
                            
                            """)
                            print(cedSecimInput)
                            if cedSecimInput=="y":
                                print("OK- ÇED seçimi tamamlandı")
                                yazinom=num
                                break
                                
                            elif cedSecimInput=="n":
                                print("Seçim yapılmadı, proje adına geri dönülecek...")
                            



                        elif "kayıt bulundu," in kayit2:
                            print( "Çok kayıt bulundu-2")
                            cedSecimInput=input("""
                            

                            *****************************
                            Projeyi listeden sen seç(y/n): 
                            *****************************
                            
                            
                            """)
                            print(cedSecimInput)
                            if cedSecimInput=="y":
                                print("OK- ÇED seçimi tamamlandı")
                                yazinom=num
                                break
                                
                            elif cedSecimInput=="n":
                                print("Seçim yapılmadı, proje adına geri dönülecek...")
                break    
                
            print("for bitti")
                    


        print("x")
        time.sleep(3) 
        
        for kunye in range(1,13):    
            #print(kunye)
            # kunyeAdı=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[1]/strong").text
            kunyeAdı=self.browser.find_element(By.XPATH,f"/html/body/div/div[2]/div[2]/table/tbody/tr[{kunye}]/td[1]").text
            self.browser.implicitly_wait(60)
            print(kunyeAdı)
            ilAdison=""
            if kunyeAdı=="İl":
                ilAdi=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[2]").text.strip(":")
                self.browser.implicitly_wait(60)
                
                
                if re.search('[,]', str(ilAdi).strip()):    
                    textList=str(ilAdi).split(",")
                    #print(textList)
                    textil=str(textList[0]).strip(":")
                    ilAdison=textil.strip()
                    #print(ilAdison)
                    break
                else:    
                    textList=str(ilAdi).strip().split(" ")
                    #print(textList)
                    textil=(textList[0]).strip(":")
                    ilAdison=textil.strip()
                    #print(ilAdison)                    
                    break
        print(ilAdison)
            
            


        for kunye in range(1,13):    
            kunyeAdı=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[1]").text
            self.browser.implicitly_wait(60)
            ilceAdison=""
            if kunyeAdı=="İlçe":
                ilceAdi=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[2]").text.strip(":")
                self.browser.implicitly_wait(60)
                if re.search('[,]', str(ilceAdi).strip()):    
                    textList=str(ilceAdi).split(",")
                    #print(textList)
                    textilce=str(textList[0]).strip(":")
                    ilceAdison=textilce.strip()
                    #print(ilceAdison)
                    break
                else:    
                    textList=str(ilceAdi).strip().split(" ")
                    #print(textList)
                    textilce=str(textList[0]).strip(":")
                    ilceAdison=textilce.strip()
                    #print(ilceAdison)
                    break
        print(ilceAdison)    
            
        

        ruhsatNoNoson=""    
        for kunye in range(1,13):    
            kunyeAdı=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[1]").text
            self.browser.implicitly_wait(60)
            if kunyeAdı=="Ruhsat Numarası":
                
                ruhsatNo=self.browser.find_element(By.XPATH,f"//*[@id='main']/table/tbody/tr[{kunye}]/td[2]").text.strip(":")
                self.browser.implicitly_wait(60)
                #print(ruhsatNo)
                ruhsatNom=re.findall(r'\d+', ruhsatNo)

                for i in ruhsatNom:
                    if int(i)<2150:
                        ruhsatNom.remove(i) 
                        #print(ruhsatNom)

                    if len(ruhsatNom)==0:                    
                        ruhsatNoNoson=ilAdison
                    elif len(ruhsatNom)!=0:
                        ruhsatNoNo=str(ruhsatNom[0]) 
                        ruhsatNoNoson=ruhsatNoNo   #çedin ruhsat nosunu alma 
                    break               
                
                #print(ruhsatNoNoson) 
                break
                    

            elif kunyeAdı!="Ruhsat Numarası":
                if yazinom!="":
                    ruhsatNoNoson=yazinom
                else:
                    ruhsatNoNoson=ilAdison
        print(ruhsatNoNoson)

    

# ÇED KLASÖRÜ İÇİNDE KLASÖR OLUŞTURMA

        dosyaListem=os.listdir("C:/Users/tugbacanan.oguz/Desktop/CED")    #ÇED kaydedilen klasörün listesini alıyor, yeni klasör numarasını tespit etmek için
        
        cedListe=[]
        cedListeNom=""
        for dosya in dosyaListem:
            cedNo=int((dosya.split("-"))[0])
            cedListe.append(cedNo)
            cedListe.sort()
            cedListeNo=str(cedListe[-1]+1)
            cedListeNom=cedListeNo
        #print(cedListe)
        #print(cedListeNom)
        print(cedListeNom)
        print(ruhsatNoNoson)
        print(ilAdison)
        print(ilceAdison)



        if len(listemson)==0:           #yazının konusundaki numaraların atıldığı listenin boş olması halinde  
            newFileName=((cedListeNom)+(f"-{ruhsatNoNoson}_"+f"{ilAdison}_"+f"{ilceAdison}"))     #yeni klasörün adını oluşturma
            #newFileName=newFileNames
            print(newFileName)            

                    
        elif len(listemson)!=0:    #yazının konusundaki numaraların atıldığı listenin dolu olması halinde 
            for num in listemson:  
                print(num)

                if int(num)<2500:   #yazının konusundaki numaraların atıldığı listedeki numaraların 2500'den küçük olması halinde proje adı inputuyla kullanıcıya sorar
                    newFileName=((cedListeNom)+(f"-{ruhsatNoNoson}_"+f"{ilAdison}_"+f"{ilceAdison}"))     #yeni klasörün adını oluşturma
                    #newFileName=newFileNames
                    print(newFileName)                                              
                            
                else:
                    newFileName=((cedListeNom)+(f"-{num}_"+f"{ilAdison}_"+f"{ilceAdison}")) 
                break
                    
        import glob
        os.mkdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}')
        x="x"
        excel=ExcelDosya((int(cedListeNom)+2),1,cedListeNom.title())
        excel.sıraYaz()
        excel=ExcelDosya((int(cedListeNom)+2),3,ilAdison.title())
        excel.ililcetarihYaz()
        excel=ExcelDosya((int(cedListeNom)+2),4,ilceAdison.title())
        excel.ililcetarihYaz()
        excel=ExcelDosya((int(cedListeNom)+2),10, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNom)+2),11, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNom)+2),12, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNom)+2),13, x)
        excel.xYaz()


       


# EÇED KML KAYDETME
        import urllib.request
        import glob
        #import requests 

        self.browser.find_element(By.XPATH,"//*[@id='home']/li[1]/a").click()   # Konum linkine tıklama
        self.browser.implicitly_wait(60)
        time.sleep(2)

        self.browser.find_element(By.XPATH,"//*[@id='main']/div[2]/a").click()   #konum sayfasından "tıklayınız" url sini tıklama
        self.browser.implicitly_wait(60)
        time.sleep(5)


        while True:
            list_of_files_kml = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
            latest_file_kml = max(list_of_files_kml, key=os.path.getctime)
            print(latest_file_kml)
            dosyaAdikml=str(latest_file_kml).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
            if dosyaAdikml.endswith("crdownload"):
                print("daha yüklenmedi, yüklemeye devam edilecek...")
                time.sleep(5)
            else:
                print("yükleme tamamlandı")
                break 


        kml_path=f"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{dosyaAdikml}"
        shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdikml}", kml_path)   # kml yi, indiği Downloads klasöründen ilgili klasöre taşır 
        print("dosya taşındı")                                                       # !! dosya adında Manisa-Spil Ã\x87BD(eksiklikleri tamamlanan).pdf'deki gibi değişik karakterler olması halinde 
                                                                                     #kullanılacak bir try except metoduyla FileNotFoundError hatası halinde başka bir dosya adı kullandır.

# #ÇED/BAŞVURU/PTD DOSYALARINI İNDİRME

        liNum=self.browser.find_elements(By.XPATH,"/html/body/div/div[2]/div[1]/div/ul[1]/li")
        self.browser.implicitly_wait(60)
        time.sleep(2)
        print(len(liNum))


    #OKUNMAYAN PDF İŞLEMLERİ
        if pdfokunabilirmi=="okunabilirpdf":
        
            print("okunabilir pdf, işlemlere devam ediliyor")
        elif pdfokunabilirmi=="okunmayanpdf":
            print("***OKUNAMAYAN PDF***")
        
            cedmi=input("""

            ******************************
            ÇED mi?(y/n):
            *******************************

            """)
            if cedmi=="n":
                print("işlemlere devam ediliyor...")
            elif cedmi=="y":
                for i in range(1,(int(len(liNum))+2)):
                    dagitimlimi=self.browser.find_element(By.XPATH,f"/html/body/div/div[2]/div[1]/div/ul[1]/li[4]").text  
                    self.browser.implicitly_wait(60)
                    print(i)

                    if dagitimlimi=="Dağıtımlı Yazıyı Gör" or "HKT Dağıtımlı Yazıyı Gör":
                        self.browser.find_element(By.XPATH,f"/html/body/div/div[2]/div[1]/div/ul[1]/li[4]").click()
                                                            
                        self.browser.implicitly_wait(60)
                        print("dağıtımlıyı gördü")
                        break
                    else:
                        print("dağıtımlı yazısını görmüyor")
                    
                time.sleep(1)
                while True:
                    list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                    latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                    print(latest_file_rapor)
                    dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                    if dosyaAdi.endswith("crdownload"):
                        print("daha yüklenmedi, yüklemeye devam edilecek...")
                        time.sleep(10)
                    else:
                        print("yükleme tamamlandı")
                        break 

                cedyazipdfindir=f"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{dosyaAdi}"
                shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", cedyazipdfindir)



                    
        cedTip=""
         
        if len(liNum)<4:
            self.browser.implicitly_wait(60)
            self.browser.find_element(By.XPATH,"//*[@id='expList']/li").click() #proje dosyaları açılan listeyi tıklıyor
            self.browser.implicitly_wait(60)
            self.browser.find_element(By.XPATH,"//*[@id='expList']/li/ul/li[1]/a").click() # raporlar satırını tıklıyor
            self.browser.implicitly_wait(60)

            kayitSayi=self.browser.find_element(By.XPATH,"//*[@id='main']/form[2]/span[1]").text    #kaç tane kayıt olduğunu text olarak getiriyor
            self.browser.implicitly_wait(60)
            kayitSayiNo=re.findall(r'\d+', kayitSayi)    # sayıları extract ediyor
            kayitSayiNom=int(kayitSayiNo[0])
            print(kayitSayiNom)

            cedORptd=self.browser.find_element(By.XPATH,f"//*[@id='dosya']").text   #Çed Başvuru mu PTD mi tespit eder
            self.browser.implicitly_wait(60)
            print(cedORptd)    
            row=1
            ptdListem=[]
            while row<=kayitSayiNom:
                x="Proje Tanıtım Dosyası"
                y="ÇED Başvuru Dosyası"
                ptd=self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{row}]/td[2]").text
                self.browser.implicitly_wait(60)
            
                ptdListem.append(ptd) 
                row+=1
            print(ptdListem) 
            if x in ptdListem:
                indexPtd=((ptdListem.index(x))+1)
                print(f"{indexPtd}PTD")
                self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{indexPtd}]/td[4]/a").click()
                self.browser.implicitly_wait(60)
                time.sleep(10)

                cedTip="PTD"
                
            elif y in ptdListem:
                indexCbd=(ptdListem.index(y))+1
                print(f"{indexCbd}ÇBD")
                self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{indexCbd}]/td[4]/a").click()  
                self.browser.implicitly_wait(60)
                time.sleep(10)
                cedTip="ÇBD"

            import glob    
            import time
            time.sleep(1)

            while True:
                list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                print(latest_file_rapor)
                dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                if dosyaAdi.endswith("crdownload"):
                    print("daha yüklenmedi, yüklemeye devam edilecek...")
                    time.sleep(10)
                else:
                    print("yükleme tamamlandı")
                    break 



            from pathlib import Path
            sizeofRep=Path(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}").stat().st_size/1024
            
            print(sizeofRep)

            ptd_cbd_path=f"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{dosyaAdi}"
            shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", ptd_cbd_path)   # ptd/çbd raporunu, indiği Downloads klasöründen ilgili klasöre taşır 
                                                                                # !! dosya adında Manisa-Spil Ã\x87BD(eksiklikleri tamamlanan).pdf'deki gibi değişik karakterler olması halinde 
                                                                                #kullanılacak bir try except metoduyla FileNotFoundError hatası halinde başka bir dosya adı kullandır.
            time.sleep(2)
            print("OK-dosya indirme")

        elif len(liNum)>=4:
            ifCed=self.browser.find_element(By.XPATH,"//*[@id='home']/li[4]/a").text   # sayfada ÇED RAPORU GÖR sekmesi varsa ÇED Raporu prosedürü uygulanır 
            self.browser.implicitly_wait(60)
        
            if ifCed=="ÇED RAPORU GÖR":
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='home']/li[4]/a").click()
                self.browser.implicitly_wait(60)
                print("Çed Raporu")

                self.browser.find_element(By.XPATH,"//*[@id='surec']/tbody/tr[1]/td[7]/input").click()    # ilk satırdaki gör butonunu tıklar
                self.browser.implicitly_wait(60)
                        
                #linkCed=self.browser.find_element(By.XPATH,"//*[@id='projeSurecForm']/table/tbody/tr[2]/td/a").get_attribute("href") # ced'in indirileceği linki alır
                self.browser.find_element(By.XPATH,"//*[@id='projeSurecForm']/table/tbody/tr[2]/td/a").click()
                self.browser.implicitly_wait(60)

                time.sleep(1)
                newCedFile=self.browser.find_element(By.XPATH,"//*[@id='projeSurecForm']/table/tbody/tr[2]/td/a/font").text #dosyanın orjinal adını alır
                self.browser.implicitly_wait(60)
                

                from pathlib import Path
                time.sleep(1)
 
                cedTip="ÇED Raporu"

                while True:
                    list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                    latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                    print(latest_file_rapor)
                    dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                    if dosyaAdi.endswith("crdownload"):
                        print("daha yüklenmedi, yüklemeye devam edilecek...")
                        time.sleep(15)
                    else:
                        print("yükleme tamamlandı")
                        break 

                ced_path=f"C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/{newCedFile}"
                shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", ced_path)
                time.sleep(1)
                sizeofRep=Path(rf"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{newCedFile}").stat().st_size/1024
                print(sizeofRep)

                                
            else:
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='expList']/li").click() #proje dosyaları açılan listeyi tıklıyor
                self.browser.implicitly_wait(60)
                self.browser.find_element(By.XPATH,"//*[@id='expList']/li/ul/li[1]/a").click() # raporlar satırını tıklıyor
                self.browser.implicitly_wait(60)

                kayitSayi=self.browser.find_element(By.XPATH,"//*[@id='main']/form[2]/span[1]").text    #kaç tane kayıt olduğunu text olarak getiriyor
                self.browser.implicitly_wait(60)
                kayitSayiNo=re.findall(r'\d+', kayitSayi)    # sayıları extract ediyor
                kayitSayiNom=int(kayitSayiNo[0])
                print(kayitSayiNom)

                cedORptd=self.browser.find_element(By.XPATH,f"//*[@id='dosya']").text   #Çed Başvuru mu PTD mi tespit eder
                self.browser.implicitly_wait(60)
                print(cedORptd)    
                row=1
                ptdListem=[]
                while row<=kayitSayiNom:
                    x="Proje Tanıtım Dosyası"
                    y="ÇED Başvuru Dosyası"
                    ptd=self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{row}]/td[2]").text
                    self.browser.implicitly_wait(60)
                
                    ptdListem.append(ptd) 
                    row+=1
                print(ptdListem) 
                if x in ptdListem:
                    indexPtd=((ptdListem.index(x))+1)
                    print(f"{indexPtd}PTD")
                    self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{indexPtd}]/td[4]/a").click()
                    self.browser.implicitly_wait(60)
                    time.sleep(10)
                    cedTip="PTD"
                    #dosyaAdi=self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{indexPtd}]/td[1]").text        # dosyaAdi ced sistemindeki indirilecek dosyanın adını kaydeder (aşağıdaki list_of_files satırı ile başlayan kod bunun yerine geçti)
                    #print(dosyaAdi)
                    
                elif y in ptdListem:
                    indexCbd=(ptdListem.index(y))+1
                    print(f"{indexCbd}ÇBD")
                    self.browser.find_element(By.XPATH,f"//*[@id='dosya']/tbody/tr[{indexCbd}]/td[4]/a").click()  
                    self.browser.implicitly_wait(60)
                    time.sleep(10)
                    cedTip="ÇBD"
                  
                    
                import glob    
                import time
                time.sleep(1)

                
                

                while True:
                    list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                    latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                    print(latest_file_rapor)
                    dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                    if dosyaAdi.endswith("crdownload"):
                        print("daha yüklenmedi, yüklemeye devam edilecek...")
                        time.sleep(10)
                    else:
                        print("yükleme tamamlandı")
                        break         

                    
                from pathlib import Path
                sizeofRep=Path(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}").stat().st_size/1024
                
                print(sizeofRep)

                ptd_cbd_path=f"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{dosyaAdi}"
                shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", ptd_cbd_path)   # ptd/çbd raporunu, indiği Downloads klasöründen ilgili klasöre taşır 
                                                                                    # !! dosya adında Manisa-Spil Ã\x87BD(eksiklikleri tamamlanan).pdf'deki gibi değişik karakterler olması halinde 
                                                                                    #kullanılacak bir try except metoduyla FileNotFoundError hatası halinde başka bir dosya adı kullandır.
                time.sleep(2)
                print("OK-dosya indirme2")

        print(cedTip)

        self.browser.switch_to.window(self.browser.window_handles[0]) 
        self.browser.implicitly_wait(60)
        sizeWindow=self.browser.get_window_size()
        self.browser.implicitly_wait(60) 
        self.browser.switch_to.frame("ustYaziOnizlemeId1")
        self.browser.implicitly_wait(60)
        self.browser.maximize_window()
        self.browser.implicitly_wait(60)
        self.browser.find_element(By.CSS_SELECTOR,"#download").click() # pencere max iken "download" tıklar üst yazıyı indirir
        self.browser.implicitly_wait(60)
        #self.browser.find_element(By.XPATH,"//*[@id='secondaryDownload']").click()  # açılan menüden "indir" seçeneğini tıklayarak yazıyı indirir                    
        self.browser.implicitly_wait(60)
        time.sleep(3)
        list_of_files_ust = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
        latest_file_ust = max(list_of_files_ust, key=os.path.getctime)
        print(latest_file_ust)
        dosyaAdi=str(latest_file_ust).lstrip(f"C:/Users/tugbacanan.oguz/Downloads/")
        ustYazipath=f"C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf"
        shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", ustYazipath) 
        sizeWindowwidth=sizeWindow["width"]
        sizeWindowheight=sizeWindow["height"]
        time.sleep(2)  
        self.browser.set_window_size(sizeWindowwidth,sizeWindowheight)
        time.sleep(2)  
        self.browser.implicitly_wait(60)  
        print(sizeWindow)
        print("size-OK")
        self.browser.implicitly_wait(60)            
        time.sleep(3)   
    
        

# BELGENETE GERİ DÖN

        print("yeni sekme")

        self.browser.switch_to.window(self.browser.window_handles[0])   # belgenet sekmesine geri dön yeni yazı yazmak için
    
        self.browser.implicitly_wait(60)
    
        
#YENİ EVRAK
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//span[contains(text(),'Evrak İşlemleri')]"))).click()  #evrak işlemleri
        
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//a[@id='topMenuForm2:ust:0:ust:0:ust:0:ust']"))).click()   #evrak oluştur
        self.browser.implicitly_wait(60)
        time.sleep(4)

#İLGİLERİ SEKMESİ
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:leftTab:uiRepeat:3:cmdbutton']").click() 
        self.browser.implicitly_wait(60)
        time.sleep(4)
        self.browser.find_element(By.XPATH,"//div[@id='window1']//li[3]").click() 
        self.browser.implicitly_wait(60)
        time.sleep(2)
        self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:evrakAramaText']").send_keys(sayiYazi)  #sayiText ile tespit edilen sayı no alınacak(1)
        self.browser.implicitly_wait(60)
        time.sleep(2)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:dokumanAraButton']").click()  #doküman ara butonuna tıklar  
        self.browser.implicitly_wait(60)        
        time.sleep(2)

        ilgiKayit=self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[3]/div/div/div/div[1]/table/tbody/tr/td").text
        self.browser.implicitly_wait(60)
        
        if ilgiKayit=="Listelenecek Veri Bulunamamıştır.":  #eğer ilgi yazı sayısı sistemde bulunamazsa metin olarak girer
            self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/ul/li[1]/a").click()   #Dosya Ekle sekmesine tıklar
            self.browser.implicitly_wait(60)
            time.sleep(2)
            self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:fileUploadButtonA_input']").send_keys(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf')        #dosya pathini yüklemek üzere gönderir
            self.browser.implicitly_wait(60)
            time.sleep(5)           
            text=self.browser.find_element(By.CSS_SELECTOR,"#yeniGidenEvrakForm\:ilgiIslemleriTabView\:dosyaEkleTab > table:nth-child(2) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(3)").text
            self.browser.implicitly_wait(60)
            print(text)
            
            
            if text in str(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf'):
                WebDriverWait(self.browser,120).until(EC.text_to_be_present_in_element((By.ID, "yeniGidenEvrakForm:ilgiIslemleriTabView:dosyaAdi"),"ustyazi.pdf"))

                self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:dosyaAciklama']").send_keys(f"{tarihYazi} tarihli ve {sayiYazim} sayılı yazı.")
                self.browser.implicitly_wait(60)
                time.sleep(1)
                self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[1]/div/button[1]/span[2]").click()  #ekle butonuna tıkla
                self.browser.implicitly_wait(60)
                time.sleep(3)
                            
            else:
                time.sleep(1)

            print("İLGİ METİN OLARAK EKLENDİ")
        else:
            self.browser.find_element(By.XPATH,"//span[@class='lobibox-close']").click()  #işlem başarılı notificationını kapatma
            self.browser.implicitly_wait(60)
            time.sleep(4)
            self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c document-follow']").click() # ilgi ekle butonuna(+) tıklar
            self.browser.implicitly_wait(60)   
        print("OK-ilgileri")
        time.sleep(4)

        
#EKLERİ SEKMESİ
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kullaniciEkleri']").click()     #Ekleri sekmesine tıklar
        self.browser.implicitly_wait(60)
    #KML UPLOAD
        #kmlfilePath=(rf"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{ruhsatNoNoson}.kml")   ???   # diğer dosyadan kml'nin pathi alınarak kullanılacak (2)
        
        self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:evrakEkTabView:fileUploadButtonA_input']").send_keys(kml_path)
        self.browser.implicitly_wait(60)
        time.sleep(5)

        
        while True:
            text=self.browser.find_element(By.CSS_SELECTOR,"label[id='yeniGidenEvrakForm:evrakEkTabView:dosyaAdi']").text
            self.browser.implicitly_wait(60)
            kml_path=konuCharHarf.char(kml_path)
            print(text)  
            print(kml_path)  
            if text in str(kml_path):
                self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("Konum")
                self.browser.implicitly_wait(60)
                time.sleep(1)
                self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaEkleButton']").click()
                self.browser.implicitly_wait(60)
                print("OK-kml")
                break              
            else:
                time.sleep(1)
                print("kml yüklenemedi")
        time.sleep(4)  

        linkBulut=""
    #ÇED UPLOAD
        if cedTip=="ÇED Raporu":
            if sizeofRep>=99000: 
                print(sizeofRep) 
                print("ÇED Raporu Tarım Buluta yüklenecek")
                self.browser.execute_script("window.open()")
                self.browser.implicitly_wait(60)
                self.browser.switch_to.window(self.browser.window_handles[2])   # tarım bulut sayfasını yeni sekmede açar
                self.browser.implicitly_wait(60)
                self.browser.get("https://dosya.tarimorman.gov.tr/app/tr-TR/App/Transfer/TarimBulut") 
                self.browser.implicitly_wait(60)
                time.sleep(2)
                self.browser.find_element(By.XPATH,"//*[@id='cerezPolitikasiDivRemove']").click()  # çerezleri kapatma (X) butonuna tıklar
                self.browser.implicitly_wait(60)
                time.sleep(3)
            

                file_input=self.browser.find_element(By.ID,"dosyaInput")
                self.browser.implicitly_wait(60)
                #time.sleep(60)
                time.sleep(2)
                #os.chdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}')
                file_input.send_keys(rf"C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/{newCedFile}") # oluşturulan fake input'un içine upload edilecek pdf dosyasını gönderir
                self.browser.implicitly_wait(60)
                time.sleep(2)
                self.browser.find_element(By.XPATH,"/html/body/div[1]/div/div[1]/div/div/div/div[2]/form/div[3]/button[1]").click()  # link bağlantısı alın butonuna tıklar
                self.browser.implicitly_wait(60)
                print("OK4")

                

                self.browser.find_element(By.XPATH,"//button[normalize-space()='Yükle']").click()   #yükle butonuna tıklar
                self.browser.implicitly_wait(60)
                WebDriverWait(self.browser,3000).until(EC.text_to_be_present_in_element((By.XPATH, "//span[normalize-space()='Linki Kopyala']"),("Linki Kopyala")))   # linki kopyala butonu görünene(dosya upload'u tamamlanana) kadar bekler

                self.browser.find_element(By.XPATH,"//*[@id='transferPanel']/div[2]/div[5]/button").click()  #linki kopyala butonunu tıklar, linki kopyalar
                self.browser.implicitly_wait(60)
                import pyperclip
                linkBulut=pyperclip.paste() #kopyalanan linki "linkBulut" objesi içine yapıştırır

                print(linkBulut)
                self.browser.switch_to.window(self.browser.window_handles[0])   #belgenet ekleri sekmesine geri döner
                self.browser.implicitly_wait(60)
            
            elif sizeofRep<=99000:

        #ÇED UPLOAD
                

                time.sleep(3)
                self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:evrakEkTabView:fileUploadButtonA_input']").send_keys(fr"C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/{newCedFile}")
                self.browser.implicitly_wait(60)
                time.sleep(5)             
                
                WebDriverWait(self.browser,120).until(EC.text_to_be_present_in_element((By.ID, "yeniGidenEvrakForm:evrakEkTabView:dosyaAdi"),(str(dosyaAdi).lstrip(dosyaAdi))+".pdf"))
                time.sleep(5) 
                print("OK-dosya")
                if ifCed=="ÇED RAPORU GÖR":

                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("ÇED Raporu") #ÇBD, ÇED ya da PTD olmasına
                    self.browser.implicitly_wait(60)
                
                else: 
                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("Dosya") #ÇBD, ÇED ya da PTD olmasına
                    self.browser.implicitly_wait(60)
                time.sleep(3)

                self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaEkleButton']").click()             
                self.browser.implicitly_wait(60)          

            
                print("OK-ekleri") 


        
        else:   
            if sizeofRep<=99000: 
                print(sizeofRep)         

        #PTD-ÇBD UPLOAD
                

                time.sleep(3)
                self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:evrakEkTabView:fileUploadButtonA_input']").send_keys(ptd_cbd_path)
                self.browser.implicitly_wait(60)
                time.sleep(5)             
                
                WebDriverWait(self.browser,120).until(EC.text_to_be_present_in_element((By.ID, "yeniGidenEvrakForm:evrakEkTabView:dosyaAdi"),(str(dosyaAdi).lstrip(dosyaAdi))+".pdf"))
                time.sleep(5) 
                print("OK-dosya")
                if "Proje Tanıtım Dosyası" in cedORptd:

                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("PTD") #ÇBD, ÇED ya da PTD olmasına
                    self.browser.implicitly_wait(60)
                elif "ÇED Başvuru Dosyası" in cedORptd:

                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("ÇBD") #ÇBD, ÇED ya da PTD olmasına
                    self.browser.implicitly_wait(60)
                else: 
                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("Dosya") #ÇBD, ÇED ya da PTD olmasına
                    self.browser.implicitly_wait(60)
                time.sleep(3)

                self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaEkleButton']").click()   
                self.browser.implicitly_wait(60)                    

            
                print("OK-ekleri")    

            elif sizeofRep>=99000: 
                print(sizeofRep)
                print("PTD/ÇBD Tarım Buluta yüklenecek")
                self.browser.execute_script("window.open()")
                self.browser.implicitly_wait(60)
                self.browser.switch_to.window(self.browser.window_handles[2])   # tarım bulut sayfasını yeni sekmede açar
                self.browser.implicitly_wait(60)
                self.browser.get("https://dosya.tarimorman.gov.tr/app/tr-TR/App/Transfer/TarimBulut") 
                self.browser.implicitly_wait(60)
                time.sleep(2)
                self.browser.find_element(By.XPATH,"//*[@id='cerezPolitikasiDivRemove']").click()  # çerezleri kapatma (X) butonuna tıklar
                self.browser.implicitly_wait(60)
                time.sleep(5)

                # self.browser.find_element(By.XPATH,"//button[@class='btn btn-secondary btn-elevate btn-icon btnTransferGorunum']").click()  # link bağlantısı alın butonuna tıklar
                # print("OK1")
                #time.sleep(3)

                # # GEÇİCİ/FAKE INPUT OLUŞTURMA
                # JS_DROP_FILE = """
                #     var target = arguments[0],
                #         offsetX = arguments[1],
                #         offsetY = arguments[2],
                #         document = target.ownerDocument || document,
                #         window = document.defaultView || window;

                #     var input = document.createElement('INPUT');
                #     input.type = 'file';
                #     input.onchange = function () {
                #     var rect = target.getBoundingClientRect(),
                #         x = rect.left + (offsetX || (rect.width >> 1)),
                #         y = rect.top + (offsetY || (rect.height >> 1)),
                #         dataTransfer = { files: this.files };

                #     ['dragenter', 'dragover', 'drop'].forEach(function (name) {
                #         var evt = document.createEvent('MouseEvent');
                #         evt.initMouseEvent(name, !0, !0, window, 0, 0, 0, x, y, !1, !1, !1, !1, 0, null);
                #         evt.dataTransfer = dataTransfer;
                #         target.dispatchEvent(evt);
                #     });

                #     setTimeout(function () { document.body.removeChild(input); }, 25);
                #     };
                #     document.body.appendChild(input);
                #     return input;
                # """

                # drop_target=self.browser.find_element(By.ID,'transferDosya') # visible olan herhangi bir elementi bul
                # self.browser.implicitly_wait(60)
                time.sleep(2)
                file_input=self.browser.find_element(By.ID,"dosyaInput")
                self.browser.implicitly_wait(60)
                time.sleep(2)

                file_input.send_keys(ptd_cbd_path) # oluşturulan fake input'un içine upload edilecek pdf dosyasını gönderir
                self.browser.implicitly_wait(60)
                time.sleep(2)
                print("OK4")
                self.browser.find_element(By.XPATH,"/html/body/div[1]/div/div[1]/div/div/div/div[2]/form/div[3]/button[1]").click()  # link bağlantısı alın butonuna tıklar
                self.browser.implicitly_wait(60)
                print("OKlinkbuton") 

                

                self.browser.find_element(By.XPATH,"//button[normalize-space()='Yükle']").click()   #yükle butonuna tıklar
                self.browser.implicitly_wait(60)                
                time.sleep(2)
                WebDriverWait(self.browser,250).until(EC.text_to_be_present_in_element((By.XPATH, "//span[normalize-space()='Linki Kopyala']"),("Linki Kopyala")))   # linki kopyala butonu görünene(dosya upload'u tamamlanana) kadar bekler
                self.browser.implicitly_wait(60) 
                time.sleep(1)
                self.browser.find_element(By.XPATH,"/html/body/div[1]/div/div[1]/div/div/div/div[2]/div[5]/button").click()  #linki kopyala butonunu tıklar, linki kopyalar
                self.browser.implicitly_wait(60) 
                
                import pyperclip
                linkBulut=pyperclip.paste() #kopyalanan linki "linkBulut" objesi içine yapıştırır

                print(linkBulut)

                self.browser.switch_to.window(self.browser.window_handles[0])   #belgenet ekleri sekmesine geri döner
                self.browser.implicitly_wait(60) 

        #EDİTÖR SEKMESİ
    
        time.sleep(4)     
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click() #editör sekmesini tıklar
        self.browser.implicitly_wait(60)
        time.sleep(8) 
        textIlgi=self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//span").text   #ilginin textini alır
        self.browser.implicitly_wait(60)
        print(textIlgi)

        yazisiRep=textIlgi.replace("yazısı.", "yazı.")
        textlist=yazisiRep.split()
        if "tarihli" in textlist:
            indexProjesi=(textlist.index("tarihli"))-1
        nihaiIlgi=" ".join(textlist[(indexProjesi):])
        print(nihaiIlgi)        #ilginin textinin düzeltilmiş nihai hali

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//span").click()     #ilginin üzerini tıklayarak text box'ın açılmasını sağlar
        self.browser.implicitly_wait(60)

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//textarea").send_keys(nihaiIlgi)    #açılan text boxa düzeltilmiş ilgi textini gönderir
        self.browser.implicitly_wait(60)

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//tbody/tr/td/span[@class='ui-inplace ui-hidden-container']/span[@class='ui-inplace-content ui-inputwrapper-filled']/span[@class='ui-inplace-editor']/button[1]").click()   # tik butonuna basarak yeni ilgiyi kaydeder
        self.browser.implicitly_wait(60)

        time.sleep(2)
            
        #self.browser.switch_to.frame(self.browser.find_elements_by_tag_name("iframe")[0])  #yazı yazılan frame'e geçer
        print("OKİ-Editör")
        print("dOKİ")

        
        #BİLGİLERİ SEKMESİ

        time.sleep(3)     
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kullaniciBilgileri']").click() #bilgileri sekmesini tıklar
        self.browser.implicitly_wait(60)
        time.sleep(3)     
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:21:eklenecekKlasorlerLov:favoriteTreeButton']").click() #kaldırılacak klasör
        self.browser.implicitly_wait(60)
        time.sleep(3)
        self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:evrakBilgileriList:21:eklenecekKlasorlerLov:lovTree:0']/span/span[3]/div/span[1]").click()    #klasör adı seç    
        self.browser.implicitly_wait(60)
        time.sleep(4)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:20:akisLov:treeButton']").click()   #onay akışı yıldız butonu tıkla
        self.browser.implicitly_wait(60)
        time.sleep(3)
        self.browser.find_element(By.XPATH,"//li[@id='yeniGidenEvrakForm:evrakBilgileriList:20:akisLov:lovTree:0']//span[@role='treeitem']//span[@role='treeitem']//div[contains(@class,'expandCollapseLovItem')]").click() #daire içi seçeneğini tıkla

        self.browser.implicitly_wait(60)
        time.sleep(3)

        # #VEKALET OLMASI HALİNDE 
        # vekalet= self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:akistaVekaletliVarPanelDialog_title']").is_displayed()
        # print(vekalet)

        # if vekalet==True:
        #     self.browser.find_element(By.XPATH,f"//div[@class='ui-radiobutton-box ui-widget ui-corner-all ui-state-default']").click()
        #     self.browser.implicitly_wait(60)
        #     self.browser.find_element(By.XPATH,f"//button[@id='yeniGidenEvrakForm:akistaVekaletliVarKaydetButton']").click()  
        #     print("OK-vek") 
        # else:
        #     print("vekalet yok")

        time.sleep(3)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:favoriteTreeButton']").click() #gereği yıldız butonunu tıkla
        self.browser.implicitly_wait(60)
        time.sleep(3)
        for i in range(0,4):
            
            self.browser.find_element(By.XPATH,f"//li[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:lovTree:{i}']").click()   #tüm daireleri seç
            self.browser.implicitly_wait(60)
            time.sleep(2)

        time.sleep(3)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:lovTreePanelKapat']").click()  # x butonu ile kapat
        self.browser.implicitly_wait(120)
        time.sleep(3)
        print("OK-Bilgileri")



        self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakBilgileriList:3:konuTextArea']").clear()
        self.browser.implicitly_wait(60)
        time.sleep(3)
        
        self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakBilgileriList:3:konuTextArea']").send_keys(joinKonuSon)
        self.browser.implicitly_wait(60)
        
        time.sleep(3)
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click()        #Editör sekmesine tıklar
        self.browser.implicitly_wait(60)
        time.sleep(3)
        print("ok-konu")
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kaydet']").click()        #KAYDET'e tıklar
        self.browser.implicitly_wait(200)
        time.sleep(50)
        print("OK-KAYDET")
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click() #editör sekmesini tıklar
        self.browser.implicitly_wait(200)
        time.sleep(10)

        from datetime import datetime
        import locale
        locale.setlocale(locale.LC_ALL, '') # local dili ayarlar, böylece aşağıdaki datetime çıktılarını türkçe verir

        simdi= datetime.today()
        result=datetime.strftime(simdi, "%d.%m.%Y %A") # gün.ay.yıl gün adı formatında bugünün tarihini verir
        print(result)        

        from datetime import timedelta
        for day in range(10,14):
            result1=simdi+timedelta(days=day)
            yaziGun=datetime.strftime(result1, "%A")
            if yaziGun=="Cumartesi" or yaziGun=="Pazar":
                print ("haftasonu")
                pass
            elif yaziGun!="Cumartesi" or yaziGun!="Pazar":
                yaziTarihi=datetime.strftime(result1, "%d.%m.%Y %A") # gün.ay.yıl gün adı formatında 7 gün sonrasının tarihini verir
                yaziTarihi2=datetime.strftime(result1, "%d.%m.%Y") # gün.ay.yıl formatında 7 gün sonrasının tarihini verir
                print(yaziTarihi)
                time.sleep(3)
                break

        
        import docx
        if cedTip=="ÇED Raporu": 
        
            os.chdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}')
            time.sleep(2)
            doc=docx.Document()  

            #document = Document(f"{newFileName}.docx")
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            
            time.sleep(2) 

            if sizeofRep>=99000:
                    
                yaziIcerigi=f"""
                İlgi yazıda {nihaiYaziIc} ÇED Raporuna ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                Söz konusu ÇED Raporu, "dosya indirme bağlantısı"{linkBulut} üzerinden temin edilebilir. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.

                """
            elif sizeofRep<=99000:
                    
                yaziIcerigi=f"""
                İlgi yazıda {nihaiYaziIc} ÇED Raporuna ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                Söz konusu ÇED Raporu ekte yer almaktadır. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.

                """



            yaziIcerigi2=f"""
            {geldYer1} 
            {geldYer2}            
            {geldigiYer}  
            
            Sayi:{sayiYazi}                                                                        Tarih:{tarihYazi}
            Konu: {joinKonuSon}
            
            
            İlgi yazıda {nihaiYaziIc} ÇED Raporuna ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
            Söz konusu ÇED Raporu Bakanlığımız Su Yönetimi Genel Müdürlüğü görev, yetki ve sorumlulukları çerçevesinde incelenmiş olup ekte yer alan görüşlerimiz e-ÇED sistemine yüklenmiştir. Bilgilerinizi ve gereğini arz ederim.
            """
            time.sleep(1)
            doc.add_paragraph(yaziIcerigi)
            time.sleep(1)
            doc.save(f"{newFileName}.docx")

            time.sleep(1)

            doc=docx.Document() 
            #document = Document(f"{newFileName}-gidecekyazi.docx")
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman' 
            time.sleep(1)
            doc.add_paragraph(yaziIcerigi2)
            time.sleep(1)
            doc.save(f"{newFileName}-gidecekyazi.docx")     
     
            
#EK-SYGM GÖRÜŞ
            document = Document('C:/Users/tugbacanan.oguz/Desktop/SYGM Görüş_.docx')
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            paragraph=document.add_paragraph()
            projeBold=paragraph.add_run("Proje: ")
            projeBold.bold=True
            paragraph.add_run(nihaiYaziIc)
            paragraph.alignment =3


            paragraph=document.add_paragraph(f"{geldYer1} {geldYer2}nün {tarihYazi} tarihli ve {sayiYazi} sayılı yazı ile e-ÇED sisteminde yer aldığı belirtilen ÇED Raporu tarafımızca incelenerek oluşturulan görüşler aşağıda yer almaktadır.")
            paragraph.alignment =3
            paragraph_format = document.styles['Normal'].paragraph_format
            paragraph_format.line_spacing = Pt(18)

            table=document.tables[0]
            cell = table.cell(0, 0)

            paragraph=cell.paragraphs[0]
            paragraph.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.').bold = True


            width250=Pt(250)
            section = document.sections[0]
            footer = section.footer
            tablefoot = footer.add_table(rows=1, cols=1, width=width250)
            tablefoot.alignment = WD_TABLE_ALIGNMENT.LEFT
            cellfoot=tablefoot.cell(0,0)   
            paraFoot=cellfoot.paragraphs[0]

            font_styles = document.styles
            font_charstyle = font_styles.add_style('footerStyle', WD_STYLE_TYPE.CHARACTER)
            font_object = font_charstyle.font
            font_object.size = Pt(8)
            font_object.name = 'Times New Roman'

            paraFoot.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.', style='footerStyle')

            print(paragraph.text)
            document.save(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/SYGM Görüş_.docx')

        elif (cedTip=="PTD") or (cedTip=="ÇBD"):
            if "Proje Tanıtım Dosyası" in cedORptd:
                projeTanitim="proje tanıtım dosyası"               

                os.chdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}')
                time.sleep(1)
                doc=docx.Document()            
                #document = Document(f"{newFileName}.docx")
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                
                time.sleep(1)

                if sizeofRep<=99000:
                    print(sizeofRep)
                    
                    yaziIcerigi=f"""
                    İlgi yazıda {nihaiYaziIc} {projeTanitim}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                    Söz konusu {projeTanitim} ekte yer almaktadır. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.
                    """

                elif sizeofRep>=99000:
                    print(sizeofRep)
                    yaziIcerigi=f"""
                    İlgi yazıda {nihaiYaziIc} {projeTanitim}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                    Söz konusu {projeTanitim}, "dosya indirme bağlantısı"{linkBulut} üzerinden temin edilebilir. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.
                    """
                yaziIcerigi2=f"""
                {geldYer1} 
                {geldYer2}                
                {geldigiYer}   
                
                Sayi:{sayiYazi}                                                                       Tarih:{tarihYazi}
                Konu: {joinKonuSon}


                İlgi yazıda {nihaiYaziIc} {projeTanitim}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                Söz konusu {projeTanitim} Bakanlığımız Su Yönetimi Genel Müdürlüğü görev, yetki ve sorumlulukları çerçevesinde incelenmiş olup ekte yer alan görüşlerimiz e-ÇED sistemine yüklenmiştir. Bilgilerinizi ve gereğini arz ederim.
                """
                doc.add_paragraph(yaziIcerigi)
                doc.save(f"{newFileName}.docx")


                doc=docx.Document() 
                #document = Document(f"{newFileName}-gidecekyazi.docx")
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)                
                time.sleep(1)
                doc.add_paragraph(yaziIcerigi2)
                time.sleep(1)
                doc.save(f"{newFileName}-gidecekyazi.docx")


                document = Document('C:/Users/tugbacanan.oguz/Desktop/SYGM Görüş_.docx')
                style = document.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                paragraph=document.add_paragraph()
                projeBold=paragraph.add_run("Proje: ")
                projeBold.bold=True
                paragraph.add_run(nihaiYaziIc)
                paragraph.alignment =3


                paragraph=document.add_paragraph(f"{geldYer1} {geldYer2}nün {tarihYazi} tarihli ve {sayiYazi} sayılı yazı ile e-ÇED sisteminde yer aldığı belirtilen {projeTanitim} tarafımızca incelenerek oluşturulan görüşler aşağıda yer almaktadır.")
                paragraph.alignment =3
                paragraph_format = document.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(18)

                table=document.tables[0]
                cell = table.cell(0, 0)

                paragraph=cell.paragraphs[0]
                paragraph.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.').bold = True


                width270=Pt(270)
                section = document.sections[0]
                footer = section.footer
                tablefoot = footer.add_table(rows=1, cols=1, width=width270)
                tablefoot.alignment = WD_TABLE_ALIGNMENT.LEFT
                cellfoot=tablefoot.cell(0,0)   
                paraFoot=cellfoot.paragraphs[0]

                font_styles = document.styles
                font_charstyle = font_styles.add_style('footerStyle', WD_STYLE_TYPE.CHARACTER)
                font_object = font_charstyle.font
                font_object.size = Pt(8)
                font_object.name = 'Times New Roman'

                paraFoot.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.', style='footerStyle')

                print(paragraph.text)
                document.save(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/SYGM Görüş_.docx')

            elif "ÇED Başvuru Dosyası" in cedORptd:
                cedBas="ÇED başvuru dosyası"

                os.chdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}')
                doc=docx.Document()  



                if sizeofRep<=99000:         
                    yaziIcerigi=f"""
                    İlgi yazıda {nihaiYaziIc} {cedBas}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                    Söz konusu {cedBas} ekte yer almaktadır. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.
                    """
                elif sizeofRep>=99000:
                    yaziIcerigi=f"""
                    İlgi yazıda {nihaiYaziIc} {cedBas}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
                    Söz konusu {cedBas}, "dosya indirme bağlantısı"{linkBulut} üzerinden temin edilebilir. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.
                    """

                yaziIcerigi2=f"""
                {geldYer1} 
                {geldYer2} 
                {geldigiYer} 

                Sayi:{sayiYazi}                                                                          Tarih:{tarihYazi}
                Konu: {joinKonuSon}

                
                İlgi yazıda {nihaiYaziIc} {cedBas}na ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir. 
                Söz konusu {cedBas} Bakanlığımız Su Yönetimi Genel Müdürlüğü görev, yetki ve sorumlulukları çerçevesinde incelenmiş olup ekte yer alan görüşlerimiz e-ÇED sistemine yüklenmiştir. Bilgilerinizi ve gereğini arz ederim.
                """                 
                doc.add_paragraph(yaziIcerigi)
                doc.save(f"{newFileName}.docx")
                #document = Document(f"{newFileName}.docx")
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                doc.save(f"{newFileName}.docx")
                time.sleep(1)
                doc=docx.Document() 
                time.sleep(1)
                doc.add_paragraph(yaziIcerigi2)
                time.sleep(1)
                doc.save(f"{newFileName}-gidecekyazi.docx")
                #document = Document(f"{newFileName}-gidecekyazi.docx")
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                doc.save(f"{newFileName}-gidecekyazi.docx")

                document = Document('C:/Users/tugbacanan.oguz/Desktop/SYGM Görüş_.docx')
                style = document.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                paragraph=document.add_paragraph()
                projeBold=paragraph.add_run("Proje: ")
                projeBold.bold=True
                paragraph.add_run(nihaiYaziIc)
                paragraph.alignment =3


                paragraph=document.add_paragraph(f"{geldYer1} {geldYer2}nün {tarihYazi} tarihli ve {sayiYazi} sayılı yazı ile e-ÇED sisteminde yer aldığı belirtilen {cedBas} tarafımızca incelenerek oluşturulan görüşler aşağıda yer almaktadır.")
                paragraph.alignment =3
                paragraph_format = document.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(18)

                table=document.tables[0]
                cell = table.cell(0, 0)

                paragraph=cell.paragraphs[0]
                paragraph.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.').bold = True


                width270=Pt(270)
                section = document.sections[0]
                footer = section.footer
                tablefoot = footer.add_table(rows=1, cols=1, width=width270)
                tablefoot.alignment = WD_TABLE_ALIGNMENT.LEFT
                cellfoot=tablefoot.cell(0,0)   
                paraFoot=cellfoot.paragraphs[0]

                font_styles = document.styles
                font_charstyle = font_styles.add_style('footerStyle', WD_STYLE_TYPE.CHARACTER)
                font_object = font_charstyle.font
                font_object.size = Pt(8)
                font_object.name = 'Times New Roman'

                paraFoot.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.', style='footerStyle')

                print(paragraph.text)
                document.save(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/SYGM Görüş_.docx')

        os.startfile(f'C:\\Users\\tugbacanan.oguz\\Desktop\\CED\\{(newFileName)}\\{newFileName}.docx')

        excel=ExcelDosya((int(cedListeNom)+2),14,yaziTarihi2)
        excel.ililcetarihYaz()


belge = Belgenet(username,password)
belge.signIn()
belge.letter()

