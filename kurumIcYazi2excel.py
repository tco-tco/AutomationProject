
from belgenetUser import username, password
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.select import Select
from docx import Document

from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from CED_TCO import ExcelDosya
from selenium.common.exceptions import NoSuchElementException
import konuCharHarf


import os
import shutil
import time
import glob



class Belgenet:
    def __init__(self, username, password):
        self.browser=webdriver.Chrome("C:/Users/tugbacanan.oguz/Desktop/python_temelleri_iş/PROJECT/chromedriver.exe")
        self.username=username
        self.password=password
        
             
    
    def signIn(self):
        self.browser.get("https://e-belge.tarim.gov.tr/edys-web/sistemeGiris.xhtml")
        self.browser.implicitly_wait(60)
        print("OK")
        self.browser.find_element(By.XPATH,"//*[@id='parolaSertifikaAccordion:uForm:txtUKullaniciAdi']").send_keys(self.username)
        self.browser.implicitly_wait(60)
        self.browser.find_element(By.XPATH,"//*[@id='loginUSifre']").send_keys(self.password)

        self.browser.implicitly_wait(60)

        self.browser.find_element(By.XPATH,"/html/body/div[5]/div[2]/div/div[2]/div/div[2]/div[2]/div/form/button/span").click()

        self.browser.implicitly_wait(60)
#KİŞİSEL ARŞİV2

        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//*[@id='leftMenuForm:kisiselArsivPanel:kisiselArsivKlasoruTab_header']"))).click() #kişisel arşiv
        
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"/html/body/div[9]/div[2]/form/div[2]/div[2]/div/ul/li/span/span[1]"))).click()  #altmenü kişisel arşiv

        #WebDriverWait(self.browser, 10).until(EC.element_to_be_clickable((By.XPATH,"//*[@id='leftMenuForm:kisiselArsivPanel:kisiselArsivKlasoruTree:0_1']"))).click()   # 2. sıradaki klasörü seçme
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.CSS_SELECTOR,"#leftMenuForm\:kisiselArsivPanel\:kisiselArsivKlasoruTree\:0_1 > span:nth-child(1) > span:nth-child(2)"))).click()   # 2. sıradaki klasörü seçme (FATMA)
        
        self.browser.implicitly_wait(60)
        time.sleep(2)
        


    def letter(self):
        import time
        konu=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[1]/td[2]/div[1]/h3").text   #yazının konusunu text halinde veriyor      
        time.sleep(2)    
        print(konu)

        konuYazim=str(konu).split()
        konuYazim.remove("Konu:")
        joinKonuSon=" ".join(konuYazim)

        #joinKonuSon=konuCharHarf.konuJoin(konu)
        konuDosyaAdi=konuCharHarf.dosyaKonu(konu)
        konuDosyaAdi2=konuCharHarf.harf(konuDosyaAdi)



        

        sayi=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[3]/td[1]/div[1]").text   #yazının sayısını text halinde veriyor 
        print(sayi)
        time.sleep(2)
        geldY=self.browser.find_element(By.XPATH,"//*[@id='mainInboxForm:inboxDataTable:0:evrakTable']/tbody/tr[2]/td[1]/div[1]").text   #yazının geldiği yeri text halinde veriyor 
        print(geldY)
        geldYerYazim=str(geldY).lstrip("Geldiği Yer:").strip()
        geldYerList=str(geldYerYazim).split("/")
        geldYerSon=str(geldYerList[0]).strip()
        print(geldYerSon)
        time.sleep(2)

        sayiYaziList=str(sayi).split("/")
        sayiYazi=(str(sayiYaziList[2]).split("-"))[-1].strip()
        print(sayiYazi)
        sayiYazim=str(sayiYaziList[2]).strip()
        print(sayiYazim)
        
        
        tarihYazi=str(sayiYaziList[1]).strip("Sayı: ").strip()
        print(tarihYazi)

# GELEN YAZI 

        WebDriverWait(self.browser, 20).until(EC.element_to_be_clickable((By.XPATH,"//*[@id='esm_246802192_emi_286053589']"))).click()      # gelen kutusunu tıklar
        self.browser.implicitly_wait(60)
        time.sleep(5)
        gelYaziNo=self.browser.find_element(By.CSS_SELECTOR,"a[id='esm_246802192_emi_286053589'] span[class='ui-menuitem-text']").text
        #gelYaziNo=self.browser.find_element(By.XPATH,"//a[@class='ui-menuitem-link ui-corner-all ui-menuitem-default ui-menuitem-unread ui-menuitem-selected ui-state-highlight']").text
        self.browser.implicitly_wait(60)
        time.sleep(2)        
        gelYaziNoList=str(gelYaziNo).lstrip("Gelen Evraklar ").split("/")
        gelYaziNom=int(gelYaziNoList[1].rstrip(")"))
        print(gelYaziNom)       
        

        sel100 = Select(self.browser.find_element(By.XPATH,"//div[@class='ui-paginator ui-paginator-top ui-widget-header']//select[@name='mainInboxForm:inboxDataTable_rppDD']"))   #sayfada görünen yazı sayısını 100'e çıkarır
        time.sleep(2)
        #select by select_by_visible_text() method
        sel100.select_by_visible_text("100")
        self.browser.implicitly_wait(120)
        time.sleep(6)
        print("OK-100")

        for i in range(0,gelYaziNom):
            time.sleep(2)
            self.browser.implicitly_wait(60)
            gelenKonu=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[1]/td[2]/div[1]/h3").text    # tüm gelen yazıların konusunu alır
            print(gelenKonu)

            self.browser.implicitly_wait(60)
            geldigiYer=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[2]/td/div").text        # tüm gelen yazıların geldiği yeri alır                                                        
            self.browser.implicitly_wait(60)
            print(geldigiYer)     
            gelenSayi=self.browser.find_element(By.XPATH,f"//*[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']/tbody/tr[4]/td/div").text        # tüm gelen yazıların tarihini alır
            self.browser.implicitly_wait(60)
            print(gelenSayi)   
            gelenSayiList=str(gelenSayi).split("/")
            print(gelenSayiList)

            if "-" in str(gelenSayiList[1]):
                gelenSayison=(str(gelenSayiList[1]).split("-"))[-1]
                
            else:
                gelenSayison=str(gelenSayiList[1]).lstrip("No: ")
            
            print(gelenSayison)               

            gelenTarih=str(gelenSayiList[0]).strip("Evrak Tarihi: ").strip()
            print(gelenTarih)
            print(i)
              

          
            if (konu == gelenKonu) and ((sayiYazi==gelenSayison) or (tarihYazi==gelenTarih)):
                self.browser.implicitly_wait(60)
                # sizeWindow=self.browser.get_window_size()
                # print(sizeWindow)
                #self.browser.maximize_window()             
                self.browser.find_element(By.XPATH,f"//table[@id='mainInboxForm:inboxDataTable:{i}:evrakTable']").click()  
                self.browser.implicitly_wait(60)
                #self.browser.find_element(By.XPATH,"//div[@id='evrakOnizlemeNotlarDialogId']//span[@class='ui-icon ui-icon-closethick']").click() # uyarı penceresi                  
                self.browser.implicitly_wait(60)            
                time.sleep(1)    
                #self.browser.set_window_size((sizeWindow.values()[0]),(sizeWindow.values()[1]))      
                self.browser.implicitly_wait(60)        
                time.sleep(1)                                  
                print("OK-gelen yazi bulundu")   
                print(i) 
                break  

    # FRAME YAZI ÖNİZLEME
        self.browser.implicitly_wait(60)
        self.browser.switch_to.frame("ustYaziOnizlemeId1")  # en sağdaki frame'e iframe id'si ile gitme
        print("OK2")
        self.browser.implicitly_wait(60)
        textim=self.browser.find_elements(By.XPATH,"//*[@id='viewer']/div[1]/div[2]/div")
        self.browser.implicitly_wait(60)
        time.sleep(5)
        print("OK3")
        #print(textim[12].text)
        self.browser.implicitly_wait(60)

        nihaiYaziIc=""      
        geldYer1=""
        geldYer2=""
        cedMetinListem=[]
        for text in textim:
            cedMetinListem.append(text.text)
        print(cedMetinListem)
        cedMetinlen=len(cedMetinListem)
        print(cedMetinlen)

        

        if int(len(cedMetinListem))<4:
            print("okunmayan pdf:(")
            
            geldYer1=""
            geldYer2=""
            


        elif int(len(cedMetinListem))>4:
            if "DAĞITIM YERLERİNE" in cedMetinListem:
                #konuIndex=(cedMetinListem.index("Konu"))+3
                yaziIcIndex=(cedMetinListem.index("DAĞITIM YERLERİNE"))+1
                print(yaziIcIndex)
                tcIndex=(cedMetinListem.index("T.C."))
                geldYer1=cedMetinListem[(int(tcIndex)+1)]
                geldYer2=cedMetinListem[(int(tcIndex)+2)]

                if cedMetinListem[yaziIcIndex] =="İlgi" or cedMetinListem[(int(yaziIcIndex)+1)]==":":
                    
                    print("yazının ilgisi var")
                    yaziIcIndexim=(cedMetinListem.index("DAĞITIM YERLERİNE"))+4
                    print(cedMetinListem[yaziIcIndex])
                    print(cedMetinListem[(int(yaziIcIndex)+1)])
                    yaziIcIndex=yaziIcIndexim
                    print(yaziIcIndex)
                    
                else:
                    print("yazının ilgisi yok")
                    #print(cedMetinListem[yaziIcIndex])

                nihaiYaziIc=""          
                yaziMetniListe=[]
                for xindex in range((yaziIcIndex+1),(cedMetinlen-yaziIcIndex)):
                    yaziMetni=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{xindex}]").text          
                    print(yaziMetni)            
                    yaziMetniListe.append(yaziMetni)
                joinYaziMetni=" ".join(yaziMetniListe)
                #print(joinYaziMetni)
                splitYaziMetni=joinYaziMetni.split()
                print(splitYaziMetni)
                

                indexProjesi=0

                if "ederim." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederim."))+1
                    indexProjesi=int(indexProje)
                elif "ederiz." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederiz."))+1
                    indexProjesi=int(indexProje)
                else:
                    indexProje=int(len(splitYaziMetni))
                    indexProjesi=indexProje

                nihaiMetin=" ".join(splitYaziMetni[:(indexProjesi)])


                nihaiYaziIc=nihaiMetin
                print(nihaiYaziIc)

            elif "SU YÖNETİMİ GENEL MÜDÜRLÜĞÜNE" in cedMetinListem:
                yaziIcIndex=(cedMetinListem.index("SU YÖNETİMİ GENEL MÜDÜRLÜĞÜNE"))+1
                print(yaziIcIndex)
                tcIndex=(cedMetinListem.index("T.C."))
                geldYer1=cedMetinListem[(int(tcIndex)+1)]
                geldYer2=cedMetinListem[(int(tcIndex)+2)]

                if cedMetinListem[yaziIcIndex] =="İlgi" and cedMetinListem[(int(yaziIcIndex)+1)]==":":
                    
                    print("yazının ilgisi var")
                    yaziIcIndexim=(cedMetinListem.index("SU YÖNETİMİ GENEL MÜDÜRLÜĞÜNE"))+4
                    print(cedMetinListem[yaziIcIndex])
                    print(cedMetinListem[(int(yaziIcIndex)+1)])
                    yaziIcIndex=yaziIcIndexim
                    print(yaziIcIndex)
                    
                else:
                    print("yazının ilgisi yok")
                    print(cedMetinListem[yaziIcIndex])
                

                    #sayiIndex=(cedMetinListem.index("Konu"))
                    # sayiText=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{sayiIndex}]").text           # aslında def(letter)'ın ilk satırındaki x gibi yazı satırından alınabilir!!!!!!
                    # print(sayiText)

                    #projesiIndex=(cedMetinListem.index("projesi"))+1
                nihaiYaziIc=""            
                yaziMetniListe=[]
                for xindex in range((yaziIcIndex+1),(cedMetinlen-yaziIcIndex)):
                    yaziMetni=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{xindex}]").text          
                    print(yaziMetni)            
                    yaziMetniListe.append(yaziMetni)
                joinYaziMetni=" ".join(yaziMetniListe)
                #print(joinYaziMetni)
                splitYaziMetni=joinYaziMetni.split()
                print(splitYaziMetni)

                indexProjesi=0


                if "ederim." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederim."))+1
                    indexProjesi=int(indexProje)
                elif "ederiz." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederiz."))+1
                    indexProjesi=int(indexProje)
                else:
                    indexProje=int(len(splitYaziMetni))
                    indexProjesi=indexProje

                nihaiMetin=" ".join(splitYaziMetni[:(indexProjesi)])


                nihaiYaziIc=nihaiMetin
                print(nihaiYaziIc)

            elif "(Su Yönetimi Genel Müdürlüğü)" in cedMetinListem:
                yaziIcIndex=(cedMetinListem.index("(Su Yönetimi Genel Müdürlüğü)"))+1
                print(yaziIcIndex)
                if "T.C." in cedMetinListem:
                    tcIndex=(cedMetinListem.index("T.C."))
                    geldYer1=cedMetinListem[(int(tcIndex)+1)]
                    geldYer2=cedMetinListem[(int(tcIndex)+2)]
                elif "TC." in cedMetinListem:
                    tcIndex=(cedMetinListem.index("TC."))
                    geldYer1=cedMetinListem[(int(tcIndex)+1)]
                    geldYer2=cedMetinListem[(int(tcIndex)+2)]

                if cedMetinListem[yaziIcIndex] =="İlgi" and cedMetinListem[(int(yaziIcIndex)+1)]==":":
                    
                    print("yazının ilgisi var")
                    yaziIcIndexim=(cedMetinListem.index("(Su Yönetimi Genel Müdürlüğü)"))+4
                    print(cedMetinListem[yaziIcIndex])
                    print(cedMetinListem[(int(yaziIcIndex)+1)])
                    yaziIcIndex=yaziIcIndexim
                    print(yaziIcIndex)
                    
                else:
                    print("yazının ilgisi yok")
                    print(cedMetinListem[yaziIcIndex])

                nihaiYaziIc=""           
                yaziMetniListe=[]
                for xindex in range((yaziIcIndex+1),(cedMetinlen-yaziIcIndex)):
                    yaziMetni=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{xindex}]").text          
                    print(yaziMetni)            
                    yaziMetniListe.append(yaziMetni)
                joinYaziMetni=" ".join(yaziMetniListe)
                #print(joinYaziMetni)
                splitYaziMetni=joinYaziMetni.split()
                print(splitYaziMetni)

                indexProjesi=0


                if "ederim." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederim."))+1
                    indexProjesi=int(indexProje)
                elif "ederiz." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederiz."))+1
                    indexProjesi=int(indexProje)
                else:
                    indexProje=int(len(splitYaziMetni))
                    indexProjesi=indexProje

                nihaiMetin=" ".join(splitYaziMetni[:(indexProjesi)])
                #print(nihaiMetin)       
                 

                nihaiYaziIc=nihaiMetin
                print(f"{nihaiYaziIc}-???DOĞRU METİN Mİ?")

            elif "TARIM VE ORMAN BAKANLIĞINA" in cedMetinListem:
                yaziIcIndex=(cedMetinListem.index("TARIM VE ORMAN BAKANLIĞINA"))+1
                print(yaziIcIndex)
                if "T.C." in cedMetinListem:
                    tcIndex=(cedMetinListem.index("T.C."))
                    geldYer1=cedMetinListem[(int(tcIndex)+1)]
                    geldYer2=cedMetinListem[(int(tcIndex)+2)]
                elif "TC." in cedMetinListem:
                    tcIndex=(cedMetinListem.index("TC."))
                    geldYer1=cedMetinListem[(int(tcIndex)+1)]
                    geldYer2=cedMetinListem[(int(tcIndex)+2)]

                if cedMetinListem[yaziIcIndex] =="İlgi" and cedMetinListem[(int(yaziIcIndex)+1)]==":":
                    
                    print("yazının ilgisi var")
                    yaziIcIndexim=(cedMetinListem.index("TARIM VE ORMAN BAKANLIĞINA"))+4
                    print(cedMetinListem[yaziIcIndex])
                    print(cedMetinListem[(int(yaziIcIndex)+1)])
                    yaziIcIndex=yaziIcIndexim
                    print(yaziIcIndex)
                    
                else:
                    print("yazının ilgisi yok")
                    print(cedMetinListem[yaziIcIndex])
                


                nihaiYaziIc=""           
                yaziMetniListe=[]
                for xindex in range((yaziIcIndex+1),(cedMetinlen-yaziIcIndex)):
                    yaziMetni=self.browser.find_element(By.XPATH,f"//*[@id='viewer']/div[1]/div[2]/div[{xindex}]").text          
                    print(yaziMetni)            
                    yaziMetniListe.append(yaziMetni)
                joinYaziMetni=" ".join(yaziMetniListe)
                #print(joinYaziMetni)
                splitYaziMetni=joinYaziMetni.split()
                print(splitYaziMetni)

                indexProjesi=0

                if "ederim." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederim."))+1
                    indexProjesi=int(indexProje)
                elif "ederiz." in splitYaziMetni:
                    indexProje=int(splitYaziMetni.index("ederiz."))+1
                    indexProjesi=int(indexProje)
                else:
                    indexProje=int(len(splitYaziMetni))
                    indexProjesi=indexProje

                nihaiMetin=" ".join(splitYaziMetni[:(indexProjesi)])
                #print(nihaiMetin)       
                 

                nihaiYaziIc=nihaiMetin
                print(f"{nihaiYaziIc}-???DOĞRU METİN Mİ?")

            
        print("OKEYİM-konu")
        print(nihaiYaziIc)

        dosyaListem=os.listdir("C:/Users/tugbacanan.oguz/Desktop/CED/")    #ÇED kaydedilen klasörün listesini alıyor, yeni klasör numarasını tespit etmek için
        cedListe=[]
        for dosya in dosyaListem:
            cedNo=int((dosya.split("-"))[0])
            cedListe.append(cedNo)
            cedListe.sort()
            cedListeNo=cedListe[-1]+1
        #print(cedListe)
        print(cedListeNo)
        newFileName=(str(cedListeNo)+f"-sayı_{sayiYazi}_"+f"{konuDosyaAdi}")    #yeni klasörün adını oluşturma
        print(newFileName)
        os.mkdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}')

        x="x"
        excel=ExcelDosya((int(cedListeNo)+2),1,str(cedListeNo).title())
        excel.sıraYaz()
        excel=ExcelDosya((int(cedListeNo)+2),2,geldYer1.title())    # geldYer1-bakanlık, valilik vs.
        excel.ililcetarihYaz()
        excel=ExcelDosya((int(cedListeNo)+2),3,geldYer2.title())    # geldYer2-??? vs.
        excel.ililcetarihYaz()
        excel=ExcelDosya((int(cedListeNo)+2),6,konuDosyaAdi.title())    # konu
        excel.ililcetarihYaz()
        excel=ExcelDosya((int(cedListeNo)+2),10, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNo)+2),11, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNo)+2),12, x)
        excel.xYaz()
        excel=ExcelDosya((int(cedListeNo)+2),13, x)
        excel.xYaz()

        
        time.sleep(3)
        print("OK-klasör oluşt")
        
        time.sleep(1)
        self.browser.execute_script("window.open()")
        self.browser.implicitly_wait(60)
        self.browser.switch_to.window(self.browser.window_handles[1])   # yeni sekmede belgenet
        self.browser.implicitly_wait(60)
        self.browser.get("https://e-belge.tarim.gov.tr/edys-web/sistemeGiris.xhtml") 
        self.browser.implicitly_wait(60)
        self.browser.switch_to.window(self.browser.window_handles[0])
        self.browser.implicitly_wait(60)
        sizeWindow=self.browser.get_window_size()
        self.browser.implicitly_wait(60)
        self.browser.maximize_window()
        self.browser.implicitly_wait(60)
        time.sleep(5)
        self.browser.implicitly_wait(60) 
        self.browser.switch_to.frame("ustYaziOnizlemeId1")
        self.browser.implicitly_wait(60)        
        self.browser.find_element(By.CSS_SELECTOR,"#download").click() # pencere max iken "download" tıklar üst yazıyı indirir
        time.sleep(4)
        self.browser.implicitly_wait(60)
        list_of_files_ust = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
        latest_file_ust = max(list_of_files_ust, key=os.path.getctime)
        print(latest_file_ust)
        dosyaAdi=str(latest_file_ust).lstrip(f"C:/Users/tugbacanan.oguz/Downloads/")
        ustYazipath=f"C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf"
        shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi}", ustYazipath) 
        time.sleep(4)    
        print("ekler sekmesi")
        
        self.browser.switch_to.window(self.browser.window_handles[0])   # belgenet sekmesine geri dön yeni yazı yazmak için
    
        self.browser.implicitly_wait(60)

        liSayi=self.browser.find_elements(By.XPATH,"/html/body/div[10]/div[3]/div/form[2]/div/div[2]/div/div/ul/li")

        time.sleep(5)
        
        print(liSayi)
        print(len(liSayi))
        ekbosmu=""
        for i in range(1,(int(len(liSayi))+1)):

            x=self.browser.find_element(By.XPATH,f"/html/body/div[10]/div[3]/div/form[2]/div/div[2]/div/div/ul/li[{i}]/a").text   
 
            self.browser.implicitly_wait(60)
            time.sleep(1)
            print(x)
            ek=""
            if x=="Evrak Ekleri":
                self.browser.find_element(By.XPATH,"/html[1]/body[1]/div[10]/div[3]/div[1]/form[2]/div[1]/div[2]/div[1]/div[1]/ul[1]/li[5]/a[1]").click()
                os.mkdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/EK')    #ekler için yeni klasör oluşturur
                self.browser.implicitly_wait(60)
                ekler=self.browser.find_elements(By.XPATH,"//div[@class='ui-datatable ui-widget myLovGrid']//table[@role='grid']//tbody//tr")   # kaç tane ek olduğunun sayısını verir
                ek=ekler
                print(len(ek))
                time.sleep(1)

                
                for i in range(1,len(ek)+1):
                    boyutEk=self.browser.find_element(By.XPATH,f"/html/body/div[10]/div[3]/div/form[2]/div/div[2]/div/div/div/div[5]/div/div[1]/div[1]/table/tbody/tr[{i}]/td[2]").text
                    indirilebilirMi=self.browser.find_element(By.XPATH,f"/html/body/div[10]/div[3]/div/form[2]/div/div[2]/div/div/div/div[5]/div/div[1]/div[1]/table/tbody/tr[{i}]/td[5]/button").get_attribute("aria-disabled")
                    print(indirilebilirMi)
                    self.browser.implicitly_wait(60)
                    time.sleep(1)
                    if len(ek)>1:
                        #if boyutEk !=  "FİZİKSEL":
                        if indirilebilirMi == "false":
                            
                            dosyaText=self.browser.find_element(By.XPATH,f"//div[@class='ui-datatable ui-widget myLovGrid']//table[@role='grid']//tbody//tr[{i}]/td[4]/div[1]/span[1]").text
                            self.browser.implicitly_wait(60)
                            self.browser.find_element(By.XPATH,f"//div[@class='ui-datatable ui-widget myLovGrid']//table[@role='grid']//tbody//tr[{i}]/td[5]/button[1]/span[1]").click()
                            self.browser.implicitly_wait(60)
                            time.sleep(1)
                            
                            while True:
                                list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                                latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                                print(latest_file_rapor)
                                dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")

                                if dosyaAdi.endswith("crdownload"):
                                    print("daha yüklenmedi, yüklemeye devam edilecek...")
                                    time.sleep(3)
                                else:
                                    print("yükleme tamamlandı")
                                    break 
                            
                            print(dosyaText)  
                            list_of_files2 = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                            latest_file2 = max(list_of_files2, key=os.path.getctime)
                            print(latest_file2)
                            time.sleep(2)
                            dosyaAdi2=str(latest_file2).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                            #os.rename(f"C://Users/t.oguz/Downloads/{dosyaAdi}.pdf", f"I:/CED/{(newFileName)}/{dosyaAdi}.pdf")
                            #shutil.move(f"C:/Users/t.oguz/Downloads/{dosyaAdi}", f"E:/CED/{(newFileName)}/{dosyaAdi}") #İŞ DOWNLOAD KLASÖRÜ
                            ekPath=f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/EK/{dosyaAdi2}'   
                            time.sleep(4)
                            shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi2}", ekPath)       
                            time.sleep(4)  
                                
                            shutil.make_archive(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/EK', 'zip', f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/EK')
                            print(f"OK EK-{i}")
                            time.sleep(6)


                        elif indirilebilirMi == "true":
                            print(f"OK-EK-{i} BOŞ-FİZİKİ EKLERE BAK") 
                            ekbosmu=f"OK-EK-{i} BOŞ-FİZİKİ EKLERE BAK"
                    elif len(ek)<=1:

                        if indirilebilirMi == "false":
                            
                            dosyaText=self.browser.find_element(By.XPATH,f"//div[@class='ui-datatable ui-widget myLovGrid']//table[@role='grid']//tbody//tr[{i}]/td[4]/div[1]/span[1]").text
                            self.browser.implicitly_wait(60)
                            self.browser.find_element(By.XPATH,f"//div[@class='ui-datatable ui-widget myLovGrid']//table[@role='grid']//tbody//tr[{i}]/td[5]/button[1]/span[1]").click()
                            self.browser.implicitly_wait(60)
                            time.sleep(1)
                            while True:
                                list_of_files_rapor = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                                latest_file_rapor = max(list_of_files_rapor, key=os.path.getctime)
                                print(latest_file_rapor)
                                dosyaAdi=str(latest_file_rapor).lstrip("C:/Users/tugbacanan.oguz/Downloads/")
                                if dosyaAdi.endswith("crdownload"):
                                    print("daha yüklenmedi, yüklemeye devam edilecek...")
                                    time.sleep(3)
                                else:
                                    print("yükleme tamamlandı")
                                    break 
                            
                            self.browser.implicitly_wait(60)
                            
                            print(dosyaText)  
                            list_of_files2 = glob.glob('C:/Users/tugbacanan.oguz/Downloads/*') # * means all if need specific format then *.csv
                            latest_file2 = max(list_of_files2, key=os.path.getctime)
                            print(latest_file2)
                            time.sleep(2)
                            dosyaAdi2=str(latest_file2).lstrip("C:/Users/tugbacanan.oguz/Downloads/")

                            ekPath=f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/EK/{dosyaAdi2}'   
                            time.sleep(4)
                            shutil.move(f"C:/Users/tugbacanan.oguz/Downloads/{dosyaAdi2}", ekPath)       
                            time.sleep(4) 
                            print(ekPath)
                            print(dosyaAdi2)

                            print("OK-Tek EK yüklendi")
                            time.sleep(6) 
 
                        elif indirilebilirMi == "true":
                            print("EKLER BOŞ-FİZİKİ EKLERE BAK") 
                            ekbosmu="EKLER BOŞ-FİZİKİ EKLERE BAK"

                break
            else:
                print("EK YOK") 
            
        print("ekler ok")

        sizeWindowwidth=sizeWindow["width"]
        sizeWindowheight=sizeWindow["height"]
        time.sleep(2)  
        self.browser.set_window_size(sizeWindowwidth,sizeWindowheight)
        time.sleep(4)  
        self.browser.implicitly_wait(60)  
        print(sizeWindow)
        print("size-OK")

        print("yeni sekme")

        self.browser.switch_to.window(self.browser.window_handles[0])   # belgenet sekmesine geri dön yeni yazı yazmak için
    
        self.browser.implicitly_wait(60)

#YENİ EVRAK
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//span[contains(text(),'Evrak İşlemleri')]"))).click()  #evrak işlemleri
        
        WebDriverWait(self.browser, 60).until(EC.element_to_be_clickable((By.XPATH,"//a[@id='topMenuForm2:ust:0:ust:0:ust:0:ust']"))).click()   #evrak oluştur
        self.browser.implicitly_wait(60)
        time.sleep(4)

        

#İLGİLERİ SEKMESİ
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:leftTab:uiRepeat:3:cmdbutton']").click() 
        self.browser.implicitly_wait(60)
        time.sleep(4)
        self.browser.find_element(By.XPATH,"//div[@id='window1']//li[3]").click() 
        self.browser.implicitly_wait(60)
        time.sleep(2)
        self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:evrakAramaText']").send_keys(sayiYazi)  #sayiText ile tespit edilen sayı no alınacak(1)
        time.sleep(2)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:dokumanAraButton']").click()  #doküman ara butonuna tıklar  
        self.browser.implicitly_wait(60)        
        time.sleep(2)

        ilgiKayit=self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[3]/div/div/div/div[1]/table/tbody/tr/td").text
        self.browser.implicitly_wait(60)
        
        if ilgiKayit=="Listelenecek Veri Bulunamamıştır.":  #eğer ilgi yazı sayısı sistemde bulunamazsa metin olarak girer
            self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/ul/li[1]/a").click()   #Dosya Ekle sekmesine tıklar
            self.browser.implicitly_wait(60)
            time.sleep(2)
            self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:fileUploadButtonA_input']").send_keys(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf')        #dosya pathini yüklemek üzere gönderir
            self.browser.implicitly_wait(60)
            time.sleep(5)           
            text=self.browser.find_element(By.CSS_SELECTOR,"#yeniGidenEvrakForm\:ilgiIslemleriTabView\:dosyaEkleTab > table:nth-child(2) > tbody:nth-child(1) > tr:nth-child(1) > td:nth-child(1) > table:nth-child(1) > tbody:nth-child(1) > tr:nth-child(2) > td:nth-child(3)").text
            print(text)
            
            
            if text in str(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/ustyazi.pdf'):
                WebDriverWait(self.browser,120).until(EC.text_to_be_present_in_element((By.ID, "yeniGidenEvrakForm:ilgiIslemleriTabView:dosyaAdi"),"ustyazi.pdf"))

                self.browser.find_element(By.XPATH,"//*[@id='yeniGidenEvrakForm:ilgiIslemleriTabView:dosyaAciklama']").send_keys(f"{tarihYazi} tarihli ve {sayiYazim} sayılı yazı.")
                self.browser.implicitly_wait(60)
                time.sleep(1)
                self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[1]/div/button[1]/span[2]").click()  #ekle butonuna tıkla
                self.browser.implicitly_wait(60)
                time.sleep(3)
                            
            else:
                time.sleep(1)

            print("İLGİ METİN OLARAK EKLENDİ")
        else:
            self.browser.find_element(By.XPATH,"//span[@class='lobibox-close']").click()  #işlem başarılı notificationını kapatma
            self.browser.implicitly_wait(60)

            time.sleep(4)
            self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c document-follow']").click() # ilgi ekle butonuna(+) tıklar
            self.browser.implicitly_wait(60)   
        print("OK-ilgileri")
        time.sleep(4)

#EKLERİ SEKMESİ
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kullaniciEkleri']").click()     #Ekleri sekmesine tıklar
        self.browser.implicitly_wait(60)
        
        ekKlasor=os.path.exists(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/EK') 
        print(ekKlasor)   
        print(ek)    
        #print("başla")
        if ekKlasor==True:
            ekKlasorFileNum=os.listdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/EK')
            print(len(ekKlasorFileNum))

            if (len(ek)>1 and (len(ekKlasorFileNum)>0)):   
         
                self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:evrakEkTabView:fileUploadButtonA_input']").send_keys(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/EK.zip')        #dosya pathini yüklemek üzere gönderir
                time.sleep(5)
                
                text=self.browser.find_element(By.CSS_SELECTOR,"label[id='yeniGidenEvrakForm:evrakEkTabView:dosyaAdi']").text
                print(text)                
                
                if text in str(f'C:/Users/tugbacanan.oguz/Desktop/CED/{newFileName}/EK.zip'):

                    WebDriverWait(self.browser,120).until(EC.text_to_be_present_in_element((By.ID, "yeniGidenEvrakForm:evrakEkTabView:dosyaAdi"),"EK.zip"))

                    self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaAciklama']").send_keys("Doküman")
                    time.sleep(1)
                    self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaEkleButton']").click()
                    print(f"OK- dosya")
                    time.sleep(3)
                                
                else:

                    time.sleep(1)
                time.sleep(6) 



            elif (len(ek)==1) and (len(ekKlasorFileNum)>0):     
                #print("elif1-başladı")           
                self.browser.find_element(By.XPATH,"//input[@id='yeniGidenEvrakForm:evrakEkTabView:fileUploadButtonA_input']").send_keys(ekPath)      #dosya pathini yüklemek üzere gönderir
                time.sleep(5)
                
                text=self.browser.find_element(By.CSS_SELECTOR,"label[id='yeniGidenEvrakForm:evrakEkTabView:dosyaAdi']").text
                self.browser.implicitly_wait(60)
                print(text)
                print(dosyaAdi2)
                backslash="\\"
                dosyaAdi3=dosyaAdi2.lstrip(backslash)
                print(dosyaAdi3)
                
                
                if text==dosyaAdi3:
                    #print("elif1-if1-başladı")
                    print("OKK")
                    WebDriverWait(self.browser,250).until(EC.text_to_be_present_in_element((By.XPATH, "/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[1]/table/tbody/tr/td[1]/table/tbody/tr[2]/td[3]/label"),dosyaAdi3))
                    print("OKK-1")    
                    self.browser.find_element(By.XPATH,"/html/body/div[11]/div/div[2]/div/form/div[2]/div[1]/div/div[1]/table/tbody/tr/td[1]/table/tbody/tr[1]/td[3]/table/tbody/tr[1]/td/textarea").send_keys("Doküman")
                    self.browser.implicitly_wait(60)
                    time.sleep(3)
                    print("OKK-2") 
                    self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakEkTabView:dosyaEkleButton']").click()
                    print(f"OK- dosya")
                    time.sleep(3)
                                
                else:
                    print("elif1-else1-başladı")
                    time.sleep(1)
                time.sleep(4) 


        elif ekKlasor==False:

            print("EKLER BOŞ-FİZİKİ EKLERE BAK")

#EDİTÖR SEKMESİ
    
        time.sleep(4)     
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click() #editör sekmesini tıklar
        time.sleep(8) 
        textIlgi=self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//span").text   #ilginin textini alır
        print(textIlgi)

        yazisiRep=textIlgi.replace("yazısı.", "yazı.")
        textlist=yazisiRep.split()
        if "tarihli" in textlist:
            indexProjesi=(textlist.index("tarihli"))-1
        nihaiIlgi=" ".join(textlist[(indexProjesi):])
        print(nihaiIlgi)        #ilginin textinin düzeltilmiş nihai hali

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//span").click()     #ilginin üzerini tıklayarak text box'ın açılmasını sağlar

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//div[@class='ui-panel ui-widget ui-widget-content ui-corner-all evrakEditorPanel']//div[5]//textarea").send_keys(nihaiIlgi)    #açılan text boxa düzeltilmiş ilgi textini gönderir

        time.sleep(2)
        self.browser.find_element(By.XPATH,"//tbody/tr/td/span[@class='ui-inplace ui-hidden-container']/span[@class='ui-inplace-content ui-inputwrapper-filled']/span[@class='ui-inplace-editor']/button[1]").click()   # tik butonuna basarak yeni ilgiyi kaydeder

        time.sleep(2)

        print("OKİ-Editör")
        print("dOKİ")

        #BİLGİLERİ SEKMESİ

        time.sleep(3)     
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kullaniciBilgileri']").click() #bilgileri sekmesini tıklar
        self.browser.implicitly_wait(60)
        time.sleep(3)     
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:21:eklenecekKlasorlerLov:favoriteTreeButton']").click() #kaldırılacak klasör

        self.browser.implicitly_wait(60)
        time.sleep(3)
        try:
            self.browser.find_element(By.XPATH,"/html/body/div[154]/div/div[2]/ul/li/span/span[3]/div/span[2]").click()    #klasör adı seç                           
            self.browser.implicitly_wait(60)
            time.sleep(2)
            hata=""
        except NoSuchElementException as NSEE:
            hata="NSEE"
        if hata=="NSEE":
            self.browser.find_element(By.XPATH,"/html/body/div[155]/div/div[2]/ul/li/span/span[3]/div/span[2]").click()    #klasör adı seç                           
            self.browser.implicitly_wait(60)
            time.sleep(2)
        else:
            print("Başka bir hata var...")



        time.sleep(3)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:favoriteTreeButton']").click() #gereği yıldız butonunu tıkla
        time.sleep(3)
        for i in range(0,4):
            
            self.browser.find_element(By.XPATH,f"//li[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:lovTree:{i}']").click()   #tüm daireleri seç
            time.sleep(2)

        time.sleep(3)
        self.browser.find_element(By.XPATH,"//button[@id='yeniGidenEvrakForm:evrakBilgileriList:17:geregiLov:lovTreePanelKapat']").click()  # x butonu ile kapat
        self.browser.implicitly_wait(60)
        time.sleep(3)
        print("OK-Bilgileri")



        self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakBilgileriList:3:konuTextArea']").clear()
        self.browser.implicitly_wait(60)
        time.sleep(3)
        
        self.browser.find_element(By.XPATH,"//textarea[@id='yeniGidenEvrakForm:evrakBilgileriList:3:konuTextArea']").send_keys(joinKonuSon)
        
        time.sleep(3)
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click()        #Editör sekmesine tıklar
        print("ok-konu")
        self.browser.implicitly_wait(60)
        time.sleep(3)
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c kaydet']").click()        #KAYDET'e tıklar
        print("OK-KAYDET")
        self.browser.implicitly_wait(120)
        time.sleep(60)
        self.browser.find_element(By.XPATH,"//span[@class='ui-button-icon-left ui-icon ui-c editor']").click() #editör sekmesini tıklar
        self.browser.implicitly_wait(120)
        time.sleep(10)

        from datetime import datetime
        import locale
        locale.setlocale(locale.LC_ALL, '') # local dili ayarlar, böylece aşağıdaki datetime çıktılarını türkçe verir

        simdi= datetime.today()
        result=datetime.strftime(simdi, "%d.%m.%Y %A") # gün.ay.yıl gün adı formatında bugünün tarihini verir
        print(result)        

        from datetime import timedelta
        for day in range(10,14):
            result1=simdi+timedelta(days=day)
            yaziGun=datetime.strftime(result1, "%A")
            if yaziGun=="Cumartesi" or yaziGun=="Pazar":
                print ("haftasonu")
                pass
            elif yaziGun!="Cumartesi" or yaziGun!="Pazar":
                yaziTarihi=datetime.strftime(result1, "%d.%m.%Y %A") # gün.ay.yıl gün adı formatında 7 gün sonrasının tarihini verir
                yaziTarihi2=datetime.strftime(result1, "%d.%m.%Y") # gün.ay.yıl formatında 7 gün sonrasının tarihini verir
                print(yaziTarihi)
                time.sleep(3)
                break

        import docx
        
        
        os.chdir(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}')
        time.sleep(2)
        doc=docx.Document()  
        time.sleep(2) 
                
        yaziIcerigi=f"""
        İlgi yazıda {nihaiYaziIc} ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
        Söz  konusu  talebe  ilişkin  dokümanlar  ekte  yer  almaktadır. Bu kapsamda hazırlanan Daire Başkanlığı görüşünüzün en geç {yaziTarihi} günü mesai bitimine kadar Daire Başkanlığımıza gönderilmesi hususunda bilgilerinizi ve gereğini arz ederim.
        
        {ekbosmu}
        """
        yaziIcerigi2=f"""
        {geldYer1} 
        {geldYer2}            
        {geldigiYer}  
        
        Sayi:{sayiYazi}                                                                        Tarih:{tarihYazi}
        Konu: {joinKonuSon}
        
        
        İlgi yazıda {nihaiYaziIc}??? ilişkin Bakanlığımız Su Yönetimi Genel Müdürlüğü görüşü talep edilmektedir.
        Söz konusu planlanan projeye ilişkin dokümanlar Bakanlığımız Su Yönetimi Genel Müdürlüğü görev, yetki ve sorumlulukları çerçevesinde incelenmiş olup görüşlerimiz ekte yer almaktadır. Bilgilerinizi ve gereğini rica ederim.
        """
        time.sleep(1)
        doc.add_paragraph(yaziIcerigi)
        time.sleep(1)
        doc.save(f"{newFileName}.docx")
        time.sleep(1)
        doc=docx.Document() 
        time.sleep(1)
        doc.add_paragraph(yaziIcerigi2)
        time.sleep(1)
        doc.save(f"{newFileName}-gidecekyazi.docx")            
        
#EK-SYGM GÖRÜŞ
        document = Document('C:/Users/tugbacanan.oguz/Desktop/SYGM Görüş_.docx')
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        paragraph=document.add_paragraph()
        projeBold=paragraph.add_run("Proje: ")
        projeBold.bold=True
        paragraph.add_run(nihaiYaziIc)
        paragraph.alignment =3


        paragraph=document.add_paragraph(f"{geldYer1} {geldYer2}nün {tarihYazi} tarihli ve {sayiYazi} sayılı yazı ile iletilen bilgi ve belgeler incelenerek oluşturulan görüşler aşağıda yer almaktadır.")
        paragraph.alignment =3
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.line_spacing = Pt(18)

        table=document.tables[0]
        cell = table.cell(0, 0)

        paragraph=cell.paragraphs[0]
        paragraph.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.').bold = True


        width300=Pt(300)
        section = document.sections[0]
        footer = section.footer
        tablefoot = footer.add_table(rows=1, cols=1, width=width300)
        tablefoot.alignment = WD_TABLE_ALIGNMENT.LEFT
        cellfoot=tablefoot.cell(0,0)   
        paraFoot=cellfoot.paragraphs[0]

        font_styles = document.styles
        font_charstyle = font_styles.add_style('footerStyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(8)
        font_object.name = 'Times New Roman'

        paraFoot.add_run(f'Su Yönetimi Genel Müdürlüğünün "{joinKonuSon}" konulu yazısının ekidir.', style='footerStyle')

        print(paragraph.text)
        document.save(f'C:/Users/tugbacanan.oguz/Desktop/CED/{(newFileName)}/SYGM Görüş_.docx')


        os.startfile(f'C:\\Users\\tugbacanan.oguz\\Desktop\\CED\\{(newFileName)}\\{newFileName}.docx')

        excel=ExcelDosya((int(cedListeNo)+2),14,yaziTarihi2)
        excel.ililcetarihYaz()



belge = Belgenet(username,password)
belge.signIn()
belge.letter()

            